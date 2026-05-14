"""
Publish Itt — Backend Server
Cleaned up: single load_dotenv, fixed duplicate /ai/analyze-tone route,
JWT auth wired in via auth router.
"""

# ── Imports ──────────────────────────────────────────────────────────────────
from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from dotenv import load_dotenv
from pathlib import Path
import logging
import os

# Load env once, from the .env file next to server.py
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

import base64
import re
import io
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional
import uuid
from datetime import datetime, timezone
from litellm import acompletion
import prompts as P
import exports as E

# Document parsing imports
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyPDF2 import PdfReader
from fpdf import FPDF
from routers.auth import auth_router, set_db as auth_set_db, get_current_user, UserOut

# ── Database ──────────────────────────────────────────────────────────────────
mongo_url = os.environ["MONGO_URL"]
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ["DB_NAME"]]

# ── Config ────────────────────────────────────────────────────────────────────
USE_MOCK_AI = os.environ.get("USE_MOCK_AI", "false").lower() == "true"
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

# ── Logging ───────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# ── App & router ──────────────────────────────────────────────────────────────
app = FastAPI(title="Publish Itt API", version="1.0.0")
api_router = APIRouter(prefix="/api")

# ============== PYDANTIC MODELS ==============


# Project Models
class ProjectBase(BaseModel):
    title: str
    user_id: Optional[str] = None  # owner — set server-side from JWT
    series_name: Optional[str] = None
    universe: Optional[str] = None
    type: Optional[str] = "novel"
    genre: Optional[str] = None
    age_group: Optional[str] = None
    writing_style: Optional[str] = None
    voice_style: Optional[str] = None
    tone_style: Optional[str] = None
    target_audience: Optional[str] = None
    pacing_preference: Optional[str] = None
    style_notes: Optional[str] = None
    status: str = "concept"
    word_count: int = 0
    summary: Optional[str] = None


class ProjectCreate(ProjectBase):
    pass


class ProjectUpdate(BaseModel):
    title: Optional[str] = None
    series_name: Optional[str] = None
    universe: Optional[str] = None
    type: Optional[str] = None
    genre: Optional[str] = None
    age_group: Optional[str] = None
    writing_style: Optional[str] = None
    voice_style: Optional[str] = None
    tone_style: Optional[str] = None
    target_audience: Optional[str] = None
    pacing_preference: Optional[str] = None
    style_notes: Optional[str] = None
    status: Optional[str] = None
    word_count: Optional[int] = None
    summary: Optional[str] = None


class Project(ProjectBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )
    updated_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )


# ============== MANUSCRIPT COLLECTION ==============
class ManuscriptBase(BaseModel):
    title: str
    raw_content: str = ""
    processed_content: str = ""
    version_id_current: Optional[str] = None
    # Extended fields from Project model for unified structure
    series_name: Optional[str] = None
    universe: Optional[str] = None
    type: str = "novel"
    status: str = "draft"
    word_count: int = 0
    summary: Optional[str] = None


class ManuscriptCreate(ManuscriptBase):
    pass


class ManuscriptUpdate(BaseModel):
    title: Optional[str] = None
    raw_content: Optional[str] = None
    processed_content: Optional[str] = None
    version_id_current: Optional[str] = None
    series_name: Optional[str] = None
    universe: Optional[str] = None
    type: Optional[str] = None
    status: Optional[str] = None
    word_count: Optional[int] = None
    summary: Optional[str] = None


class Manuscript(ManuscriptBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )
    updated_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )
    source_project_id: Optional[str] = None  # For migration tracking


# ============== VERSION COLLECTION ==============
class VersionBase(BaseModel):
    parent_type: str  # 'manuscript' or 'chapter'
    parent_id: str
    content_snapshot: str = ""
    label: str = ""
    created_by: str = ""
    user_id: Optional[str] = None  # owner — set server-side from JWT


class VersionCreate(VersionBase):
    pass


class Version(VersionBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )


# ============== NOTES COLLECTION ==============
class NoteBase(BaseModel):
    parent_type: str  # 'manuscript' or 'chapter'
    parent_id: str
    note_text: str = ""
    location_reference: str = ""
    note_type: str = ""  # 'todo', 'comment', 'revision', 'author_intent'
    user_id: Optional[str] = None  # owner — set server-side from JWT


class NoteCreate(NoteBase):
    pass


class NoteUpdate(BaseModel):
    note_text: Optional[str] = None
    location_reference: Optional[str] = None
    note_type: Optional[str] = None


class Note(NoteBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )


# ============== APPLIED IMPROVEMENTS COLLECTION ==============
class AppliedImprovementBase(BaseModel):
    action: str
    response_text: str
    project_id: Optional[str] = None
    chapter_id: Optional[str] = None
    original_content_preview: Optional[str] = None
    implemented: bool = True
    implemented_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )
    source: str = "thad_import_action"


class AppliedImprovementCreate(AppliedImprovementBase):
    pass


class AppliedImprovement(AppliedImprovementBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))


# ============== WRITING STATISTICS COLLECTION ==============
class WritingSessionBase(BaseModel):
    user_id: Optional[str] = None  # owner — set server-side from JWT
    project_id: Optional[str] = None
    chapter_id: Optional[str] = None
    date: str  # YYYY-MM-DD format
    words_added: int = 0
    words_deleted: int = 0
    time_spent_seconds: int = 0


class WritingSessionCreate(WritingSessionBase):
    pass


class WritingSession(WritingSessionBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )


class DailyStatsResponse(BaseModel):
    date: str
    total_words_added: int = 0
    total_words_deleted: int = 0
    net_words: int = 0
    total_time_seconds: int = 0
    sessions_count: int = 0
    projects_worked: List[str] = []
    chapters_worked: List[str] = []


class WritingStreakResponse(BaseModel):
    current_streak: int = 0
    longest_streak: int = 0
    last_writing_date: Optional[str] = None
    streak_dates: List[str] = []


class WritingStatsOverview(BaseModel):
    total_words_written: int = 0
    total_time_seconds: int = 0
    total_sessions: int = 0
    current_streak: int = 0
    longest_streak: int = 0
    average_words_per_day: float = 0
    average_session_minutes: float = 0
    days_active: int = 0
    weekly_words: List[dict] = []  # Last 7 days


class TodayStatsResponse(BaseModel):
    """Compact today-focused snapshot for the Dashboard momentum strip."""
    date: str
    words_today: int = 0
    time_today_seconds: int = 0
    sessions_today: int = 0
    daily_word_goal: int = 500
    goal_reached: bool = False
    current_streak: int = 0
    last_writing_date: Optional[str] = None


# Chapter Models
class ChapterBase(BaseModel):
    project_id: str
    manuscript_id: Optional[str] = None
    chapter_number: int
    title: str
    content: str = ""
    summary: str = ""
    status: str = "draft"
    version_id_current: Optional[str] = None


class ChapterCreate(ChapterBase):
    pass


class ChapterUpdate(BaseModel):
    chapter_number: Optional[int] = None
    title: Optional[str] = None
    content: Optional[str] = None
    summary: Optional[str] = None
    status: Optional[str] = None
    version_id_current: Optional[str] = None


class Chapter(ChapterBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )
    updated_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )


# ToneProfile Models
class ToneProfileBase(BaseModel):
    project_id: str
    chapter_id: Optional[str] = None
    detected_tone: str
    reading_level: str
    pacing_notes: str
    voice_notes: str
    suggestions: List[str] = []


class ToneProfileCreate(ToneProfileBase):
    pass


class ToneProfile(ToneProfileBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )


# ArtAsset Models
class ArtAssetBase(BaseModel):
    project_id: str
    chapter_id: Optional[str] = None
    type: str  # cover, chapter_header, spot_illustration
    style_preset: str
    prompt_used: str
    status: str = "generated"
    image_reference: Optional[str] = None


class ArtAssetCreate(ArtAssetBase):
    pass


class ArtAsset(ArtAssetBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )


# BookArtProfile Models - Visual identity for a manuscript
class BookArtProfileBase(BaseModel):
    project_id: str
    genre: str = ""
    age_group: str = ""
    mood: str = ""
    art_style_preferences: str = ""
    color_palette: str = ""
    reference_notes: str = ""
    ai_summary: Optional[str] = None  # AI-generated visual identity summary


class BookArtProfileCreate(BookArtProfileBase):
    pass


class BookArtProfileUpdate(BaseModel):
    genre: Optional[str] = None
    age_group: Optional[str] = None
    mood: Optional[str] = None
    art_style_preferences: Optional[str] = None
    color_palette: Optional[str] = None
    reference_notes: Optional[str] = None
    ai_summary: Optional[str] = None


class BookArtProfile(BookArtProfileBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )
    updated_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )


# StylePreset Models
class StylePresetBase(BaseModel):
    name: str
    description: str
    visual_style: str
    mood: str
    color_palette: Optional[str] = None
    user_id: Optional[str] = None  # owner — set server-side from JWT


class StylePresetCreate(StylePresetBase):
    pass


class StylePresetUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    visual_style: Optional[str] = None
    mood: Optional[str] = None
    color_palette: Optional[str] = None


class StylePreset(StylePresetBase):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    created_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )


# AI Request/Response Models
class RewriteRequest(BaseModel):
    content: str
    tone: str = "warm and engaging"
    project_id: Optional[str] = None


class SummarizeRequest(BaseModel):
    content: str


class AnalyzeToneRequest(BaseModel):
    content: str
    project_id: Optional[str] = None
    chapter_id: Optional[str] = None


class OutlineRequest(BaseModel):
    project_title: Optional[str] = None
    project_summary: str
    target_chapter_count: int = 10


class WorkflowAnalysisRequest(BaseModel):
    status_description: str


class ToneAnalysisRequest(BaseModel):
    content: str
    project_id: str
    chapter_id: Optional[str] = None
    section_info: Optional[str] = None
    intended_tone: Optional[str] = None
    goals: Optional[str] = None
    age_group: Optional[str] = None


class ToneStyleAnalysisResponse(BaseModel):
    tone_analysis: str
    style_analysis: str
    suggestions: List[str]
    reading_level: Optional[str] = None


class ArtPromptRequest(BaseModel):
    project_id: str
    chapter_id: Optional[str] = None
    style_preset: str
    prompt_type: str  # cover, chapter_header, spot_illustration
    context: str


class AskThadRequest(BaseModel):
    query: str
    context: Optional[str] = None


class AIChatRequest(BaseModel):
    message: str
    chapter_content: str = ""
    selected_text: Optional[str] = None
    project_id: Optional[str] = None


class AIResponse(BaseModel):
    response: str
    module: str


# Market Intelligence Request Models
class BookIdeasRequest(BaseModel):
    universe: str = "My Story Universe"
    count: int = 10


class MarketAnalysisRequest(BaseModel):
    genre: str = "fiction"
    age_group: Optional[str] = None


class CustomerResearchRequest(BaseModel):
    book_idea: str


class MarketOutlineRequest(BaseModel):
    book_idea: str
    chapter_count: int = 12


class ManuscriptDraftRequest(BaseModel):
    book_idea: str
    word_count: int = 30000


class BookDescriptionRequest(BaseModel):
    book_title: str
    book_summary: str


class SalesAnalysisRequest(BaseModel):
    sales_data: str


# Import Analysis Request Models
class ImportAnalysisRequest(BaseModel):
    content: str
    filename: Optional[str] = None
    project_id: Optional[str] = None
    chapter_id: Optional[str] = None


class ImportActionRequest(BaseModel):
    action: str  # autoformat, remove_notes, store_notes, convert_notes, split_chapters, lantern_path, full_qa, extract_summaries, extract_characters, extract_glossary
    content: str
    project_id: Optional[str] = None
    chapter_id: Optional[str] = None


class ImportAnalysisResponse(BaseModel):
    analysis: str
    structure_issues: List[str]
    notes_detected: List[str]
    style_issues: List[str]
    formatting_issues: List[str]
    lore_issues: List[str]
    word_count: int
    estimated_reading_level: str
    recommended_actions: List[str]


# Workflow Stage Analysis Models
class WorkflowStageAnalysisRequest(BaseModel):
    manuscript: str
    section_info: Optional[str] = None
    workflow_stage: Optional[str] = None
    goals: Optional[str] = None
    time_away: Optional[str] = None
    age_group: Optional[str] = None
    project_id: Optional[str] = None


class WorkflowStageAnalysisResponse(BaseModel):
    stage: str  # Idea Drop, Outline, Draft, Revise, Polish, Complete
    message: str
    next_steps: List[str]
    progress_percent: int


# Writing Momentum Analysis Models
class WritingMomentumRequest(BaseModel):
    daily_words: int = 0
    weekly_words: int = 0
    streak: int = 0
    total_words: int = 0
    session_minutes: int = 0
    time_away: Optional[str] = None
    goals: Optional[str] = None
    age_group: Optional[str] = None


class WritingMomentumResponse(BaseModel):
    message: str
    suggestions: List[str]


# ============== SYSTEM PROMPTS ==============
# Imported from prompts.py — see that module for the voice rules.

GLOBAL_SYSTEM_PROMPT = P.GLOBAL_SYSTEM_PROMPT
MANUSCRIPT_SYSTEM_PROMPT = P.MANUSCRIPT_SYSTEM_PROMPT
WORKFLOW_SYSTEM_PROMPT = P.WORKFLOW_SYSTEM_PROMPT
WORKFLOW_STAGE_SYSTEM_PROMPT = P.WORKFLOW_STAGE_SYSTEM_PROMPT
TONE_STYLE_SYSTEM_PROMPT = P.TONE_STYLE_SYSTEM_PROMPT
TONE_STYLE_ANALYSIS_SYSTEM_PROMPT = P.TONE_STYLE_ANALYSIS_SYSTEM_PROMPT
WRITING_MOMENTUM_SYSTEM_PROMPT = P.WRITING_MOMENTUM_SYSTEM_PROMPT
BOOK_ART_PROFILE_SYSTEM_PROMPT = P.BOOK_ART_PROFILE_SYSTEM_PROMPT
ART_STUDIO_SYSTEM_PROMPT = P.ART_STUDIO_SYSTEM_PROMPT
MARKET_INTELLIGENCE_SYSTEM_PROMPT = P.MARKET_INTELLIGENCE_SYSTEM_PROMPT
IMPORT_ANALYSIS_SYSTEM_PROMPT = P.IMPORT_ANALYSIS_SYSTEM_PROMPT


# ============== AI HELPER FUNCTIONS ==============


def _extract_text_block(user_message: str) -> str:
    if not user_message:
        return ""

    # Extract ONLY the content AFTER the LAST "Text:"
    matches = list(re.finditer(r"Text:\s*", user_message, re.IGNORECASE))

    if not matches:
        return user_message.strip()

    last_match = matches[-1]
    start = last_match.end()

    text_block = user_message[start:].strip()

    # Remove any accidental trailing sections (just in case)
    text_block = re.split(r"\n##|\n[A-Z][a-zA-Z ]+:", text_block)[0]

    return text_block.strip()


def _mock_rewrite(
    text: str,
    tone_style: str = "",
    pacing_preference: str = "",
    target_audience: str = "",
    voice_style: str = "",
    style_notes: str = "",
) -> str:
    import re

    if not text or not text.strip():
        return (
            "[THAD PROFILE: IMPROVE WRITING]\n\n"
            "## Intent\n"
            "- Strengthen clarity, readability, and flow.\n"
            "- Preserve the original meaning and emotional direction.\n\n"
            "## Suggested Revision\n"
            "No text provided.\n\n"
            "## What Changed\n"
            "- Waiting for source content.\n\n"
            "## Next Pass\n"
            "- Paste a paragraph, scene, or chapter excerpt to improve."
        )

    def split_sentences(value: str):
        value = re.sub(r"\s+", " ", value).strip()
        if not value:
            return []
        parts = re.split(r"(?<=[.!?…])\s+", value)
        return [p.strip() for p in parts if p.strip()]

    def protect_case_replace(sentence: str, old: str, new: str) -> str:
        pattern = re.compile(re.escape(old), re.IGNORECASE)

        def repl(match):
            original = match.group(0)
            if original.isupper():
                return new.upper()
            if original[:1].isupper():
                return new[:1].upper() + new[1:]
            return new

        return pattern.sub(repl, sentence)

    def simplify_for_younger(sentence: str) -> str:
        replacements = [
            ("glorious", "wonderful"),
            ("gorgeously", "very"),
            ("catastrophe", "disaster"),
            ("calamity", "mess"),
            ("dramatically", "loudly"),
            ("murmured", "said softly"),
            ("scanning", "looking over"),
            ("whispered", "said quietly"),
            ("clung", "hung"),
            ("drooped", "hung"),
        ]
        for old, new in replacements:
            sentence = protect_case_replace(sentence, old, new)
        sentence = re.sub(r"\bperhaps\b", "maybe", sentence, flags=re.IGNORECASE)
        sentence = re.sub(r"\bapproximately\b", "about", sentence, flags=re.IGNORECASE)
        return sentence

    def apply_tone(sentence: str, tone: str) -> str:
        lowered = tone.lower()

        # 🌑 Dark / eerie tone
        if any(
            word in lowered
            for word in ["dark", "eerie", "spooky", "mysterious", "atmospheric"]
        ):
            sentence = protect_case_replace(sentence, "quiet", "strangely quiet")
            sentence = protect_case_replace(sentence, "Too quiet.", "Far too quiet.")
            sentence = protect_case_replace(sentence, "empty", "strangely empty")
            sentence = protect_case_replace(sentence, "Cold.", "Cold and unnatural.")
            sentence = protect_case_replace(sentence, "Still.", "Unnaturally still.")
            sentence = protect_case_replace(
                sentence, "I don’t like this.", "Something feels wrong."
            )

        # 🌤 Warm / gentle tone
        if any(word in lowered for word in ["warm", "gentle", "hopeful", "heartfelt"]):
            sentence = protect_case_replace(sentence, "Cold.", "Cool, but not harsh.")
            sentence = protect_case_replace(
                sentence, "I don’t like this.", "This feels strange."
            )

        # 🎭 Playful tone
        if any(word in lowered for word in ["playful", "whimsical", "light"]):
            sentence = protect_case_replace(
                sentence, "gasped dramatically", "gasped with flair"
            )
            sentence = protect_case_replace(
                sentence, "I don’t like this.", "This is kinda weird..."
            )

        return sentence

    def apply_voice(sentence: str, voice: str) -> str:
        lowered = voice.lower()

        if "cinematic" in lowered:
            sentence = protect_case_replace(sentence, "The air felt", "The air seemed")
            sentence = protect_case_replace(
                sentence, "He stood, scanning", "He rose and surveyed"
            )

        if "direct" in lowered:
            sentence = protect_case_replace(
                sentence, "That’s when he saw it.", "Then he saw it."
            )
            sentence = protect_case_replace(
                sentence, "He lifted it carefully.", "He lifted it."
            )

        if "lyrical" in lowered:
            sentence = protect_case_replace(
                sentence, "The air felt", "The air seemed to feel"
            )
            sentence = protect_case_replace(
                sentence, "bare and tired", "bare, tired, and worn"
            )

        return sentence

    def apply_pacing(sentence: str, pacing: str) -> str:
        lowered = pacing.lower()

        if "fast" in lowered or "brisk" in lowered or "quick" in lowered:
            sentence = protect_case_replace(
                sentence, "the moment he stepped into", "when he entered"
            )
            sentence = protect_case_replace(
                sentence, "nearly tumbling into", "nearly crashing into"
            )
            sentence = protect_case_replace(
                sentence, "He lifted it carefully.", "He lifted it."
            )
            sentence = re.sub(r"\bthat is\b", "that's", sentence, flags=re.IGNORECASE)
            sentence = re.sub(r"\bit is\b", "it's", sentence, flags=re.IGNORECASE)

        if "slow" in lowered or "gentle" in lowered or "reflective" in lowered:
            sentence = protect_case_replace(
                sentence, "He stood, scanning", "He stood for a moment, scanning"
            )
            sentence = protect_case_replace(
                sentence,
                "He touched the soil.",
                "He reached down and touched the soil.",
            )
            sentence = protect_case_replace(
                sentence,
                "He touched the leaves.",
                "He brushed his hand over the leaves.",
            )

        return sentence

    def apply_style_notes(sentence: str, notes: str) -> str:
        lowered = notes.lower()

    def soften_sentence(sentence: str) -> str:
        sentence = sentence.replace(
            "The air felt… quiet.", "The air felt strangely quiet."
        )
        sentence = sentence.replace("Too quiet.", "Far too quiet.")
        sentence = sentence.replace("Cold.", "Cold and still.")
        sentence = sentence.replace("Still.", "Unnaturally still.")
        sentence = sentence.replace("Limp.", "Soft and limp.")
        return sentence

        if "vivid" in lowered:
            sentence = protect_case_replace(
                sentence, "bare and tired", "bare, tired, and worn"
            )
            sentence = protect_case_replace(
                sentence, "not a single berry", "not one berry"
            )

        if "easy to read" in lowered or "simple" in lowered:
            sentence = simplify_for_younger(sentence)

        return sentence

    original_sentences = split_sentences(text)
    rewritten_sentences = []

    younger_reader = any(
        phrase in (target_audience or "").lower()
        for phrase in [
            "middle grade",
            "young reader",
            "children",
            "kids",
            "elementary",
            "3rd",
            "4th",
            "5th",
        ]
    )

    changed_count = 0

    for sentence in original_sentences:
        updated = sentence

        updated = apply_tone(updated, tone_style)
        updated = apply_voice(updated, voice_style)
        updated = apply_pacing(updated, pacing_preference)

        updated = soften_sentence(updated)

        if younger_reader:
            updated = simplify_for_younger(updated)

        if style_notes.strip():
            updated = apply_style_notes(updated, style_notes)

        # Generic light cleanup so more than one sentence changes.
        updated = protect_case_replace(
            updated, "knew something was wrong", "felt that something was wrong"
        )
        updated = protect_case_replace(
            updated, "That’s when he saw it.", "Then he saw it."
        )
        updated = protect_case_replace(
            updated,
            "A small piece of bark, wedged under a stone.",
            "A small piece of bark was wedged under a stone.",
        )
        updated = protect_case_replace(
            updated, "He lifted it carefully.", "He lifted it with care."
        )
        updated = protect_case_replace(updated, "come on", "come along")

        updated = re.sub(r"\s+", " ", updated).strip()

        if updated != sentence:
            changed_count += 1

        rewritten_sentences.append(updated)

    rewritten_text = " ".join(rewritten_sentences).strip()

    if rewritten_text == text.strip() or changed_count == 0:
        rewritten_text = protect_case_replace(
            rewritten_text, "something was wrong", "something felt wrong"
        )
        rewritten_text = protect_case_replace(
            rewritten_text, "Too quiet.", "Far too quiet."
        )
        rewritten_text = protect_case_replace(
            rewritten_text, "He lifted it carefully.", "He lifted it with care."
        )

    change_notes = [
        "- Smoothed sentence flow for easier reading.",
        "- Preserved the core meaning, emotional direction, and story facts.",
    ]

    if tone_style.strip():
        change_notes.append(
            f"- Lightly shaped the wording toward a {tone_style.strip()} tone."
        )
    if pacing_preference.strip():
        change_notes.append(
            f"- Adjusted sentence rhythm to support a {pacing_preference.strip()} pace."
        )
    if target_audience.strip():
        change_notes.append(
            f"- Kept readability aligned with {target_audience.strip()}."
        )
    if voice_style.strip():
        change_notes.append(
            f"- Nudged the phrasing toward the project's {voice_style.strip()} voice."
        )
    if style_notes.strip():
        change_notes.append("- Applied the project's extra style notes where possible.")

    return (
        "[THAD PROFILE: IMPROVE WRITING]\n\n"
        "## Intent\n"
        "- Strengthen clarity, readability, and flow.\n"
        "- Preserve the original meaning and emotional direction.\n\n"
        "## Suggested Revision\n"
        f"{rewritten_text}\n\n"
        "## What Changed\n" + "\n".join(change_notes) + "\n\n## Next Pass\n"
        "- Review whether you want this to feel warmer, sharper, darker, or more playful."
    )


def _mock_summary(text: str) -> str:
    cleaned = " ".join(text.split())
    if not cleaned:
        return (
            "CHAPTER SUMMARY\n\n"
            "Quick Summary:\n"
            "- No text provided.\n\n"
            "Key Beats:\n"
            "- No opening beat detected\n"
            "- No middle development detected\n"
            "- No ending movement detected\n\n"
            "Emotional Arc:\n"
            "- Add chapter text to evaluate the emotional movement.\n\n"
            "Important Notes:\n"
            "- Waiting for source material\n"
            "- Summary workflow is ready once content is provided"
        )

    preview = cleaned[:400]
    return (
        "CHAPTER SUMMARY\n\n"
        "Quick Summary:\n"
        "- This chapter segment moves the reader through a clear piece of story progression while maintaining the manuscript's core tone and direction. The material suggests forward motion, a central focus, and enough narrative context to support downstream editing and testing.\n\n"
        "Key Beats:\n"
        "- The section opens by establishing the current focus, situation, or narrative frame\n"
        "- A meaningful development, realization, or tension beat appears in the middle\n"
        "- The ending leaves the reader with momentum into the next section\n\n"
        "Emotional Arc:\n"
        "- The emotional movement feels steady and readable, with the passage carrying a sense of progression rather than stalling in place.\n\n"
        "Important Notes:\n"
        f"- Source preview: {preview}\n"
        "- This summary is formatted to feel like a reusable THAD response profile"
    )


def _mock_outline(text: str) -> str:
    cleaned = " ".join(text.split())
    preview = cleaned[:240] if cleaned else "No source text provided."
    return (
        "STORY OUTLINE\n\n"
        "Story Goal:\n"
        "- Shape the material into a clear sequence of setup, escalation, turning point, and forward momentum so the manuscript feels intentional from section to section.\n\n"
        "Main Beats:\n"
        "1. Establish the opening situation and orient the reader to the current focus\n"
        "2. Introduce a goal, pressure point, or unresolved question\n"
        "3. Deepen the tension through action, discovery, or complication\n"
        "4. Deliver a meaningful emotional or plot turn\n"
        "5. Close with a transition that naturally leads into the next chapter or scene\n\n"
        "Next Step:\n"
        f"- Expand this outline into chapter beats using the source direction suggested here: {preview}"
    )


def _mock_structure_analysis(text: str) -> str:
    cleaned = " ".join(text.split())
    preview = cleaned[:300] if cleaned else "No text provided."
    return (
        "STRUCTURE ANALYSIS\n\n"
        "Overall Assessment:\n"
        "- The manuscript appears structurally readable, with a visible beginning, development area, and transition toward a payoff or next beat.\n\n"
        "Strengths:\n"
        "- The opening appears to establish context in a usable way\n"
        "- The middle suggests active development rather than static description\n"
        "- The material has enough shape to support revision and outlining\n\n"
        "Weaknesses:\n"
        "- Some transitions may feel softer than the surrounding material\n"
        "- Repetition could reduce forward momentum in longer passages\n"
        "- Scene or section turns may need clearer emphasis\n\n"
        "Recommendations:\n"
        "1. Tighten repeated phrasing so each paragraph advances something new\n"
        "2. Strengthen transitions between major beats or sections\n"
        "3. Clarify turning points so the reader feels the structural movement more strongly\n\n"
        f"Source Preview:\n- {preview}"
    )


async def get_ai_response(
    system_prompt: str, user_message: str, session_id: str = None
) -> str:
    if USE_MOCK_AI:
        text = user_message
        lower_prompt = user_message.lower()

        if (
            "create outline" in lower_prompt
            or "story outline" in lower_prompt
            or "outline" in lower_prompt
        ):
            return _mock_outline(text)

        if "analyze structure" in lower_prompt or "structure analysis" in lower_prompt:
            return _mock_structure_analysis(text)

        if (
            "chapter summary" in lower_prompt
            or "summarize" in lower_prompt
            or "summary" in lower_prompt
        ):
            return _mock_summary(text)

        if "rewrite" in lower_prompt:
            voice_match = re.search(r"Voice style:\s*(.+)", user_message, re.IGNORECASE)
            tone_match = re.search(r"Tone style:\s*(.+)", user_message, re.IGNORECASE)
            audience_match = re.search(
                r"Target audience:\s*(.+)", user_message, re.IGNORECASE
            )
            pacing_match = re.search(r"Pacing:\s*(.+)", user_message, re.IGNORECASE)
            style_notes_match = re.search(
                r"Style notes:\s*(.+)", user_message, re.IGNORECASE
            )

            voice_style = voice_match.group(1).strip() if voice_match else ""
            tone_style = tone_match.group(1).strip() if tone_match else ""
            target_audience = audience_match.group(1).strip() if audience_match else ""
            pacing_preference = pacing_match.group(1).strip() if pacing_match else ""
            style_notes = (
                style_notes_match.group(1).strip() if style_notes_match else ""
            )

            return _mock_rewrite(
                text,
                tone_style=tone_style,
                pacing_preference=pacing_preference,
                target_audience=target_audience,
                voice_style=voice_style,
                style_notes=style_notes,
            )

        return (
            "[MOCK AI RESPONSE]\n\n"
            "This is a mock response for development testing.\n\n"
            f"Source preview:\n{text[:800]}"
        )

    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="AI service not configured")

    if session_id is None:
        session_id = str(uuid.uuid4())

    try:
        response = await acompletion(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            api_key=OPENAI_API_KEY,
            timeout=60,
        )

        content = (
            response.choices[0].message.content
            if response and response.choices
            else None
        )
        if not content:
            raise HTTPException(status_code=500, detail="AI returned empty content")

        return content.strip()

    except Exception as e:
        logger.error(f"AI request failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI request failed: {str(e)}")


# ============== PROJECT ENDPOINTS ==============



@api_router.post("/projects", response_model=Project)
async def create_project(project: ProjectCreate, current_user: UserOut = Depends(get_current_user)):
    project_obj = Project(**project.model_dump())
    project_obj.user_id = current_user.id
    doc = project_obj.model_dump()
    await db.projects.insert_one(doc)
    return project_obj


@api_router.get("/projects", response_model=List[Project])
async def get_projects(current_user: UserOut = Depends(get_current_user)):
    try:
        projects = await db.projects.find({"user_id": current_user.id}, {"_id": 0}).to_list(1000)
        return projects
    except Exception as e:
        logger.error(f"Failed to load projects: {e}")
        return []


@api_router.get("/projects/{project_id}", response_model=Project)
async def get_project(project_id: str, current_user: UserOut = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id, "user_id": current_user.id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


@api_router.put("/projects/{project_id}", response_model=Project)
async def update_project(project_id: str, update: ProjectUpdate, current_user: UserOut = Depends(get_current_user)):
    update_data = {k: v for k, v in update.model_dump().items() if v is not None}
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()

    result = await db.projects.update_one({"id": project_id, "user_id": current_user.id}, {"$set": update_data})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Project not found")

    project = await db.projects.find_one({"id": project_id, "user_id": current_user.id}, {"_id": 0})
    return project


@api_router.delete("/projects/{project_id}")
async def delete_project(project_id: str, current_user: UserOut = Depends(get_current_user)):
    result = await db.projects.delete_one({"id": project_id, "user_id": current_user.id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Project not found")
    # Also delete related chapters
    await db.chapters.delete_many({"project_id": project_id})
    return {"message": "Project deleted successfully"}


# ============== CHAPTER ENDPOINTS ==============


@api_router.post("/chapters", response_model=Chapter)
async def create_chapter(chapter: ChapterCreate, current_user: UserOut = Depends(get_current_user)):
    # Verify the project belongs to this user
    project = await db.projects.find_one({"id": chapter.project_id, "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    chapter_obj = Chapter(**chapter.model_dump())
    doc = chapter_obj.model_dump()
    await db.chapters.insert_one(doc)

    # Update project word count
    word_count = len(chapter.content.split()) if chapter.content else 0
    project = await db.projects.find_one({"id": chapter.project_id})
    if project:
        new_word_count = project.get("word_count", 0) + word_count
        await db.projects.update_one(
            {"id": chapter.project_id},
            {
                "$set": {
                    "word_count": new_word_count,
                    "updated_at": datetime.now(timezone.utc).isoformat(),
                }
            },
        )

    return chapter_obj


@api_router.get("/chapters/project/{project_id}", response_model=List[Chapter])
async def get_chapters_by_project(project_id: str, current_user: UserOut = Depends(get_current_user)):
    # Verify project ownership
    project = await db.projects.find_one({"id": project_id, "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    chapters = (
        await db.chapters.find({"project_id": project_id}, {"_id": 0})
        .sort("chapter_number", 1)
        .to_list(1000)
    )
    return chapters


@api_router.get("/chapters/{chapter_id}", response_model=Chapter)
async def get_chapter(chapter_id: str, current_user: UserOut = Depends(get_current_user)):
    chapter = await db.chapters.find_one({"id": chapter_id}, {"_id": 0})
    if not chapter:
        raise HTTPException(status_code=404, detail="Chapter not found")
    # Verify parent project ownership
    project = await db.projects.find_one({"id": chapter["project_id"], "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Chapter not found")
    return chapter


@api_router.put("/chapters/{chapter_id}", response_model=Chapter)
async def update_chapter(chapter_id: str, update: ChapterUpdate, current_user: UserOut = Depends(get_current_user)):
    # Get old chapter for word count calculation
    old_chapter = await db.chapters.find_one({"id": chapter_id})
    if not old_chapter:
        raise HTTPException(status_code=404, detail="Chapter not found")
    # Verify parent project ownership
    project = await db.projects.find_one({"id": old_chapter["project_id"], "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Chapter not found")

    update_data = {k: v for k, v in update.model_dump().items() if v is not None}
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()

    await db.chapters.update_one({"id": chapter_id}, {"$set": update_data})

    # Update project word count if content changed
    if "content" in update_data:
        old_words = len(old_chapter.get("content", "").split())
        new_words = len(update_data["content"].split())
        word_diff = new_words - old_words

        if word_diff != 0:
            await db.projects.update_one(
                {"id": old_chapter["project_id"]},
                {
                    "$inc": {"word_count": word_diff},
                    "$set": {"updated_at": datetime.now(timezone.utc).isoformat()},
                },
            )

    chapter = await db.chapters.find_one({"id": chapter_id}, {"_id": 0})
    return chapter


@api_router.delete("/chapters/{chapter_id}")
async def delete_chapter(chapter_id: str, current_user: UserOut = Depends(get_current_user)):
    chapter = await db.chapters.find_one({"id": chapter_id})
    if not chapter:
        raise HTTPException(status_code=404, detail="Chapter not found")
    # Verify parent project ownership
    project = await db.projects.find_one({"id": chapter["project_id"], "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Chapter not found")

    # Update project word count
    word_count = len(chapter.get("content", "").split())
    await db.projects.update_one(
        {"id": chapter["project_id"]},
        {
            "$inc": {"word_count": -word_count},
            "$set": {"updated_at": datetime.now(timezone.utc).isoformat()},
        },
    )

    await db.chapters.delete_one({"id": chapter_id})
    return {"message": "Chapter deleted successfully"}


# ============== MANUSCRIPTS COLLECTION ENDPOINTS (DEPRECATED) ==============
# NOTE: These endpoints are deprecated. Use /api/projects endpoints instead.
# The projects + chapters pattern is the recommended approach.
# These endpoints will be removed in a future version.


@api_router.post("/manuscripts-collection", deprecated=True)
@api_router.get("/manuscripts-collection", deprecated=True)
@api_router.get("/manuscripts-collection/{manuscript_id}", deprecated=True)
@api_router.put("/manuscripts-collection/{manuscript_id}", deprecated=True)
@api_router.delete("/manuscripts-collection/{manuscript_id}", deprecated=True)
async def manuscripts_collection_gone(*args, **kwargs):
    """[REMOVED] This endpoint family has been removed. Use /api/projects instead."""
    raise HTTPException(status_code=410, detail="Endpoint removed. Use /api/projects instead.")


# ============== VERSIONS COLLECTION ENDPOINTS ==============


@api_router.post("/versions", response_model=Version)
async def create_version(version: VersionCreate, current_user: UserOut = Depends(get_current_user)):
    """Create a new version snapshot"""
    # Verify the parent entity belongs to this user
    owner = await _verify_parent_ownership(version.parent_type, version.parent_id, current_user.id)
    if not owner:
        raise HTTPException(status_code=404, detail="Parent not found")
    version_obj = Version(**version.model_dump())
    version_obj.user_id = current_user.id
    version_obj.created_by = current_user.id
    doc = version_obj.model_dump()
    await db.versions.insert_one(doc)
    return version_obj


@api_router.get(
    "/versions/parent/{parent_type}/{parent_id}", response_model=List[Version]
)
async def get_versions_by_parent(parent_type: str, parent_id: str, current_user: UserOut = Depends(get_current_user)):
    """Get all versions for a specific parent (manuscript or chapter)"""
    versions = (
        await db.versions.find(
            {"parent_type": parent_type, "parent_id": parent_id, "user_id": current_user.id}, {"_id": 0}
        )
        .sort("created_at", -1)
        .to_list(100)
    )
    return versions


@api_router.get("/versions/{version_id}", response_model=Version)
async def get_version(version_id: str, current_user: UserOut = Depends(get_current_user)):
    """Get a specific version by ID"""
    version = await db.versions.find_one({"id": version_id, "user_id": current_user.id}, {"_id": 0})
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    return version


@api_router.delete("/versions/{version_id}")
async def delete_version(version_id: str, current_user: UserOut = Depends(get_current_user)):
    """Delete a version"""
    result = await db.versions.delete_one({"id": version_id, "user_id": current_user.id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Version not found")
    return {"message": "Version deleted successfully"}


# ============== DATA MIGRATION ENDPOINT (DEPRECATED) ==============


@api_router.post("/migrate/projects-to-manuscripts", deprecated=True)
async def migrate_projects_to_manuscripts():
    """[DEPRECATED] Migrate existing projects to manuscripts collection.

    Note: The manuscripts_collection is deprecated. Projects + Chapters is the recommended pattern.
    This endpoint is kept for backward compatibility only.
    """
    logger.warning("Deprecated migration endpoint called.")
    projects = await db.projects.find({}, {"_id": 0}).to_list(1000)
    migrated_count = 0

    for project in projects:
        # Check if already migrated
        existing = await db.manuscripts_collection.find_one(
            {"source_project_id": project["id"]}
        )
        if existing:
            continue

        # Create manuscript from project
        manuscript = Manuscript(
            title=project.get("title", "Untitled"),
            raw_content="",
            processed_content="",
            version_id_current=None,
        )
        manuscript_doc = manuscript.model_dump()
        manuscript_doc["source_project_id"] = project["id"]  # Track migration
        manuscript_doc["series_name"] = project.get("series_name")
        manuscript_doc["universe"] = project.get("universe")
        manuscript_doc["type"] = project.get("type", "novel")
        manuscript_doc["status"] = project.get("status", "draft")
        manuscript_doc["word_count"] = project.get("word_count", 0)
        manuscript_doc["summary"] = project.get("summary")

        await db.manuscripts_collection.insert_one(manuscript_doc)

        # Update chapters to reference new manuscript
        await db.chapters.update_many(
            {"project_id": project["id"]}, {"$set": {"manuscript_id": manuscript.id}}
        )

        migrated_count += 1

    return {
        "success": True,
        "message": f"Migrated {migrated_count} projects to manuscripts",
        "total_projects": len(projects),
        "migrated": migrated_count,
    }



async def _verify_parent_ownership(parent_type: str, parent_id: str, user_id: str) -> bool:
    """Return True if the parent entity (project or chapter) belongs to user_id."""
    if parent_type in ("manuscript", "project"):
        doc = await db.projects.find_one({"id": parent_id, "user_id": user_id})
        return doc is not None
    elif parent_type == "chapter":
        chapter = await db.chapters.find_one({"id": parent_id})
        if not chapter:
            return False
        project = await db.projects.find_one({"id": chapter["project_id"], "user_id": user_id})
        return project is not None
    return False


# ============== NOTES COLLECTION ENDPOINTS ==============


@api_router.post("/notes", response_model=Note)
async def create_note(note: NoteCreate, current_user: UserOut = Depends(get_current_user)):
    """Create a new note"""
    # Verify the parent entity belongs to this user (parent_id is a project or chapter id)
    owner = await _verify_parent_ownership(note.parent_type, note.parent_id, current_user.id)
    if not owner:
        raise HTTPException(status_code=404, detail="Parent not found")
    note_obj = Note(**note.model_dump())
    note_obj.user_id = current_user.id
    doc = note_obj.model_dump()
    await db.notes.insert_one(doc)
    return note_obj


@api_router.get("/notes/parent/{parent_type}/{parent_id}", response_model=List[Note])
async def get_notes_by_parent(parent_type: str, parent_id: str, current_user: UserOut = Depends(get_current_user)):
    """Get all notes for a specific parent (manuscript or chapter)"""
    notes = (
        await db.notes.find(
            {"parent_type": parent_type, "parent_id": parent_id, "user_id": current_user.id}, {"_id": 0}
        )
        .sort("created_at", -1)
        .to_list(1000)
    )
    return notes


@api_router.get("/notes/{note_id}", response_model=Note)
async def get_note(note_id: str, current_user: UserOut = Depends(get_current_user)):
    """Get a specific note by ID"""
    note = await db.notes.find_one({"id": note_id, "user_id": current_user.id}, {"_id": 0})
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    return note


@api_router.put("/notes/{note_id}", response_model=Note)
async def update_note(note_id: str, update: NoteUpdate, current_user: UserOut = Depends(get_current_user)):
    """Update a note"""
    update_data = {k: v for k, v in update.model_dump().items() if v is not None}

    result = await db.notes.update_one({"id": note_id, "user_id": current_user.id}, {"$set": update_data})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Note not found")

    note = await db.notes.find_one({"id": note_id, "user_id": current_user.id}, {"_id": 0})
    return note


@api_router.delete("/notes/{note_id}")
async def delete_note(note_id: str, current_user: UserOut = Depends(get_current_user)):
    """Delete a note"""
    result = await db.notes.delete_one({"id": note_id, "user_id": current_user.id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Note not found")
    return {"message": "Note deleted successfully"}


# ============== WRITING STATISTICS ENDPOINTS ==============


@api_router.post("/stats/session", response_model=WritingSession)
async def log_writing_session(session: WritingSessionCreate, current_user: UserOut = Depends(get_current_user)):
    """Log a writing session"""
    session_obj = WritingSession(**session.model_dump())
    session_obj.user_id = current_user.id
    doc = session_obj.model_dump()
    await db.writing_sessions.insert_one(doc)
    return session_obj


@api_router.get("/stats/daily/{date}", response_model=DailyStatsResponse)
async def get_daily_stats(date: str, current_user: UserOut = Depends(get_current_user)):
    """Get writing stats for a specific date (YYYY-MM-DD)"""
    sessions = await db.writing_sessions.find({"user_id": current_user.id, "date": date}, {"_id": 0}).to_list(1000)

    if not sessions:
        return DailyStatsResponse(date=date)

    total_words_added = sum(s.get("words_added", 0) for s in sessions)
    total_words_deleted = sum(s.get("words_deleted", 0) for s in sessions)
    total_time = sum(s.get("time_spent_seconds", 0) for s in sessions)
    projects = list(set(s.get("project_id") for s in sessions if s.get("project_id")))
    chapters = list(set(s.get("chapter_id") for s in sessions if s.get("chapter_id")))

    return DailyStatsResponse(
        date=date,
        total_words_added=total_words_added,
        total_words_deleted=total_words_deleted,
        net_words=total_words_added - total_words_deleted,
        total_time_seconds=total_time,
        sessions_count=len(sessions),
        projects_worked=projects,
        chapters_worked=chapters,
    )


async def _compute_writing_streak(user_id: str) -> WritingStreakResponse:
    """Internal helper — compute streak for a given user_id."""
    from datetime import timedelta
    pipeline = [{"$match": {"user_id": user_id}}, {"$group": {"_id": "$date"}}, {"$sort": {"_id": -1}}]
    dates_cursor = db.writing_sessions.aggregate(pipeline)
    dates = [doc["_id"] async for doc in dates_cursor]

    if not dates:
        return WritingStreakResponse()

    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    yesterday = (datetime.now(timezone.utc) - timedelta(days=1)).strftime("%Y-%m-%d")

    current_streak = 0
    streak_dates = []

    if dates[0] == today or dates[0] == yesterday:
        current_date = datetime.strptime(dates[0], "%Y-%m-%d")
        for date_str in dates:
            date = datetime.strptime(date_str, "%Y-%m-%d")
            expected_date = current_date - timedelta(days=current_streak)
            if date == expected_date:
                current_streak += 1
                streak_dates.append(date_str)
            else:
                break

    return WritingStreakResponse(
        current_streak=current_streak,
        longest_streak=current_streak,
        last_writing_date=dates[0] if dates else None,
        streak_dates=streak_dates,
    )


@api_router.get("/stats/streak", response_model=WritingStreakResponse)
async def get_writing_streak(current_user: UserOut = Depends(get_current_user)):
    """Calculate current writing streak"""
    return await _compute_writing_streak(current_user.id)


@api_router.get("/stats/today", response_model=TodayStatsResponse)
async def get_today_stats(current_user: UserOut = Depends(get_current_user)):
    """
    Compact 'how am I doing right now' snapshot for the Dashboard momentum strip:
    today's words, today's time, current streak, and the user's daily goal.
    Single round-trip so the dashboard doesn't need three.
    """
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    sessions = await db.writing_sessions.find(
        {"user_id": current_user.id, "date": today}, {"_id": 0}
    ).to_list(1000)

    words_today = sum(s.get("words_added", 0) for s in sessions)
    time_today = sum(s.get("time_spent_seconds", 0) for s in sessions)

    streak_data = await _compute_writing_streak(current_user.id)
    goal = current_user.daily_word_goal or 500

    return TodayStatsResponse(
        date=today,
        words_today=words_today,
        time_today_seconds=time_today,
        sessions_today=len(sessions),
        daily_word_goal=goal,
        goal_reached=goal > 0 and words_today >= goal,
        current_streak=streak_data.current_streak,
        last_writing_date=streak_data.last_writing_date,
    )


@api_router.get("/stats/overview", response_model=WritingStatsOverview)
async def get_stats_overview(current_user: UserOut = Depends(get_current_user)):
    """Get overall writing statistics overview"""
    # Get all sessions for this user
    sessions = await db.writing_sessions.find({"user_id": current_user.id}, {"_id": 0}).to_list(10000)

    if not sessions:
        return WritingStatsOverview()

    total_words = sum(s.get("words_added", 0) for s in sessions)
    total_time = sum(s.get("time_spent_seconds", 0) for s in sessions)
    total_sessions = len(sessions)

    # Get unique dates
    unique_dates = list(set(s.get("date") for s in sessions))
    days_active = len(unique_dates)

    # Get streak info
    streak_data = await _compute_writing_streak(current_user.id)

    # Calculate averages
    avg_words_per_day = total_words / days_active if days_active > 0 else 0
    avg_session_minutes = (
        (total_time / total_sessions / 60) if total_sessions > 0 else 0
    )

    # Get last 7 days data
    from datetime import timedelta

    weekly_words = []
    for i in range(6, -1, -1):
        date = (datetime.now(timezone.utc) - timedelta(days=i)).strftime("%Y-%m-%d")
        day_sessions = [s for s in sessions if s.get("date") == date]
        day_words = sum(s.get("words_added", 0) for s in day_sessions)
        day_name = (datetime.now(timezone.utc) - timedelta(days=i)).strftime("%a")
        weekly_words.append({"date": date, "day": day_name, "words": day_words})

    return WritingStatsOverview(
        total_words_written=total_words,
        total_time_seconds=total_time,
        total_sessions=total_sessions,
        current_streak=streak_data.current_streak,
        longest_streak=streak_data.longest_streak,
        average_words_per_day=round(avg_words_per_day, 1),
        average_session_minutes=round(avg_session_minutes, 1),
        days_active=days_active,
        weekly_words=weekly_words,
    )


@api_router.get("/stats/weekly")
async def get_weekly_stats(current_user: UserOut = Depends(get_current_user)):
    """Get stats for the last 7 days"""
    from datetime import timedelta

    weekly_data = []
    for i in range(6, -1, -1):
        date = (datetime.now(timezone.utc) - timedelta(days=i)).strftime("%Y-%m-%d")
        day_name = (datetime.now(timezone.utc) - timedelta(days=i)).strftime("%a")

        sessions = await db.writing_sessions.find({"user_id": current_user.id, "date": date}, {"_id": 0}).to_list(
            1000
        )

        words_added = sum(s.get("words_added", 0) for s in sessions)
        time_spent = sum(s.get("time_spent_seconds", 0) for s in sessions)

        weekly_data.append(
            {
                "date": date,
                "day": day_name,
                "words": words_added,
                "time_minutes": round(time_spent / 60, 1),
                "sessions": len(sessions),
            }
        )

    return weekly_data


# ============== IMPORT MANUSCRIPT ACTION ==============


@api_router.post("/actions/import-manuscript")
async def action_import_manuscript(
    file: UploadFile = File(...), title: Optional[str] = Form(None), current_user: UserOut = Depends(get_current_user)
):
    """
    Import Manuscript Action
    Accepts a file upload and stores it in Manuscripts.raw_content
    """
    # Validate file type
    allowed_extensions = {".txt", ".docx", ".pdf", ".md"}
    filename = file.filename or "uploaded_file"
    file_ext = Path(filename).suffix.lower()

    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file_ext}. Allowed: {', '.join(allowed_extensions)}",
        )

    # Read and extract content
    content = await file.read()

    try:
        if file_ext == ".txt":
            text = extract_text_from_txt(content)
        elif file_ext == ".docx":
            text = extract_text_from_docx(content)
        elif file_ext == ".pdf":
            text = extract_text_from_pdf(content)
        elif file_ext == ".md":
            text = extract_text_from_md(content)
        else:
            text = ""
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="No text content found in the file")

    # Create manuscript record
    manuscript_title = title or filename.replace(file_ext, "")
    manuscript_obj = Manuscript(
        title=manuscript_title, raw_content=text, processed_content=""
    )
    doc = manuscript_obj.model_dump()
    await db.manuscripts_collection.insert_one(doc)

    return {
        "success": True,
        "message": f"Brought in '{manuscript_title}'.",
        "manuscript_id": manuscript_obj.id,
        "word_count": len(text.split()),
        "filename": filename,
    }


# ============== MANUSCRIPT UPLOAD ENDPOINTS ==============


def extract_text_from_txt(content: bytes) -> str:
    """Extract text from a .txt file"""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from a .docx file - handles paragraphs, tables, and headers"""
    doc = DocxDocument(io.BytesIO(content))
    text_parts = []

    # Extract from paragraphs (main body)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    # Extract from tables (some manuscripts use tables for formatting)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        text_parts.append(text)

    # Join with double newlines to preserve paragraph structure
    result = "\n\n".join(text_parts)

    # Log extraction stats for debugging
    logger.info(
        f"DOCX extraction: {len(text_parts)} text blocks, {len(result)} total chars"
    )

    return result


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from a .pdf file"""
    reader = PdfReader(io.BytesIO(content))
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return "\n\n".join(text_parts)


def extract_text_from_md(content: bytes) -> str:
    """Extract text from a .md file"""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


class UploadResponse(BaseModel):
    success: bool
    message: str
    filename: str
    content: str
    word_count: int
    chapter_id: Optional[str] = None


@api_router.post("/manuscripts/upload", response_model=UploadResponse)
async def upload_manuscript(
    file: UploadFile = File(...),
    project_id: str = Form(...),
    chapter_title: Optional[str] = Form(None),
    current_user: UserOut = Depends(get_current_user),
):
    """Upload a manuscript file and optionally create a chapter from it"""

    # Validate file type
    allowed_extensions = {".txt", ".docx", ".pdf", ".md"}
    filename = file.filename or "uploaded_file"
    file_ext = Path(filename).suffix.lower()

    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file_ext}. Allowed: {', '.join(allowed_extensions)}",
        )

    # Read file content
    content = await file.read()

    # Extract text based on file type
    try:
        if file_ext == ".txt":
            text = extract_text_from_txt(content)
        elif file_ext == ".docx":
            text = extract_text_from_docx(content)
        elif file_ext == ".pdf":
            text = extract_text_from_pdf(content)
        elif file_ext == ".md":
            text = extract_text_from_md(content)
        else:
            text = ""
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="No text content found in the file")

    # Calculate word count
    word_count = len(text.split())

    # Verify project exists
    project = await db.projects.find_one({"id": project_id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Create chapter if requested
    chapter_id = None
    if chapter_title:
        # Get next chapter number
        existing_chapters = await db.chapters.count_documents(
            {"project_id": project_id}
        )

        chapter_obj = Chapter(
            project_id=project_id,
            chapter_number=existing_chapters + 1,
            title=chapter_title,
            content=f"<p>{text.replace(chr(10), '</p><p>')}</p>",
            status="draft",
        )
        doc = chapter_obj.model_dump()
        await db.chapters.insert_one(doc)
        chapter_id = chapter_obj.id

        # Update project word count
        await db.projects.update_one(
            {"id": project_id},
            {
                "$inc": {"word_count": word_count},
                "$set": {"updated_at": datetime.now(timezone.utc).isoformat()},
            },
        )

    return UploadResponse(
        success=True,
        message=f"Read {filename}.",
        filename=filename,
        content=text,
        word_count=word_count,
        chapter_id=chapter_id,
    )


@api_router.post("/manuscripts/upload-preview")
async def preview_manuscript_upload(file: UploadFile = File(...), current_user: UserOut = Depends(get_current_user)):
    """Preview a manuscript file without creating a chapter"""

    # Validate file type
    allowed_extensions = {".txt", ".docx", ".pdf", ".md"}
    filename = file.filename or "uploaded_file"
    file_ext = Path(filename).suffix.lower()

    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file_ext}. Allowed: {', '.join(allowed_extensions)}",
        )

    # Read file content
    content = await file.read()
    logger.info(f"Upload preview: Read {len(content)} bytes from {filename}")

    # Extract text based on file type
    try:
        if file_ext == ".txt":
            text = extract_text_from_txt(content)
        elif file_ext == ".docx":
            text = extract_text_from_docx(content)
        elif file_ext == ".pdf":
            text = extract_text_from_pdf(content)
        elif file_ext == ".md":
            text = extract_text_from_md(content)
        else:
            text = ""
    except Exception as e:
        logger.error(f"Failed to parse {filename}: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="No text content found in the file")

    word_count = len(text.split())
    logger.info(
        f"Upload preview: Extracted {len(text)} chars, {word_count} words from {filename}"
    )

    # Check for chapter markers
    import re

    chapter_markers = re.findall(r"(?:CHAPTER|Chapter|chapter)\s+\d+", text)
    logger.info(
        f"Upload preview: Found {len(chapter_markers)} chapter markers: {chapter_markers[:10]}"
    )

    # Return preview (first 2000 chars)
    preview = text[:2000] + ("..." if len(text) > 2000 else "")

    return {
        "success": True,
        "filename": filename,
        "file_type": file_ext,
        "word_count": word_count,
        "preview": preview,
        "full_content": text,
    }


# ============== STYLE PRESET ENDPOINTS ==============


@api_router.post("/style-presets", response_model=StylePreset)
async def create_style_preset(preset: StylePresetCreate, current_user: UserOut = Depends(get_current_user)):
    preset_obj = StylePreset(**preset.model_dump())
    preset_obj.user_id = current_user.id
    doc = preset_obj.model_dump()
    await db.style_presets.insert_one(doc)
    return preset_obj


@api_router.get("/style-presets", response_model=List[StylePreset])
async def get_style_presets(current_user: UserOut = Depends(get_current_user)):
    presets = await db.style_presets.find({"user_id": current_user.id}, {"_id": 0}).to_list(100)
    return presets


@api_router.get("/style-presets/{preset_id}", response_model=StylePreset)
async def get_style_preset(preset_id: str, current_user: UserOut = Depends(get_current_user)):
    preset = await db.style_presets.find_one({"id": preset_id, "user_id": current_user.id}, {"_id": 0})
    if not preset:
        raise HTTPException(status_code=404, detail="Style preset not found")
    return preset


@api_router.put("/style-presets/{preset_id}", response_model=StylePreset)
async def update_style_preset(preset_id: str, update: StylePresetUpdate, current_user: UserOut = Depends(get_current_user)):
    update_data = {k: v for k, v in update.model_dump().items() if v is not None}

    result = await db.style_presets.update_one({"id": preset_id, "user_id": current_user.id}, {"$set": update_data})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Style preset not found")

    preset = await db.style_presets.find_one({"id": preset_id, "user_id": current_user.id}, {"_id": 0})
    return preset


@api_router.delete("/style-presets/{preset_id}")
async def delete_style_preset(preset_id: str, current_user: UserOut = Depends(get_current_user)):
    result = await db.style_presets.delete_one({"id": preset_id, "user_id": current_user.id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Style preset not found")
    return {"message": "Style preset deleted successfully"}


# ============== ART ASSET ENDPOINTS ==============


@api_router.post("/art-assets", response_model=ArtAsset)
async def create_art_asset(asset: ArtAssetCreate, current_user: UserOut = Depends(get_current_user)):
    project = await db.projects.find_one({"id": asset.project_id, "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    asset_obj = ArtAsset(**asset.model_dump())
    doc = asset_obj.model_dump()
    await db.art_assets.insert_one(doc)
    return asset_obj


@api_router.get("/art-assets/project/{project_id}", response_model=List[ArtAsset])
async def get_art_assets_by_project(project_id: str, current_user: UserOut = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id, "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    assets = await db.art_assets.find({"project_id": project_id}, {"_id": 0}).to_list(
        1000
    )
    return assets


@api_router.delete("/art-assets/{asset_id}")
async def delete_art_asset(asset_id: str, current_user: UserOut = Depends(get_current_user)):
    # Verify ownership via parent project
    asset = await db.art_assets.find_one({"id": asset_id})
    if not asset:
        raise HTTPException(status_code=404, detail="Art asset not found")
    project = await db.projects.find_one({"id": asset["project_id"], "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Art asset not found")
    result = await db.art_assets.delete_one({"id": asset_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Art asset not found")
    return {"message": "Art asset deleted successfully"}


# ============== BOOK ART PROFILE ENDPOINTS ==============


@api_router.post("/art-profiles", response_model=BookArtProfile)
async def create_or_update_art_profile(profile: BookArtProfileCreate, current_user: UserOut = Depends(get_current_user)):
    """Create or update a book art profile for a project"""
    project = await db.projects.find_one({"id": profile.project_id, "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    existing = await db.book_art_profiles.find_one({"project_id": profile.project_id})

    if existing:
        # Update existing profile
        update_data = profile.model_dump()
        update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
        await db.book_art_profiles.update_one(
            {"project_id": profile.project_id}, {"$set": update_data}
        )
        updated = await db.book_art_profiles.find_one(
            {"project_id": profile.project_id}, {"_id": 0}
        )
        return updated
    else:
        # Create new profile
        profile_obj = BookArtProfile(**profile.model_dump())
        doc = profile_obj.model_dump()
        await db.book_art_profiles.insert_one(doc)
        return profile_obj


@api_router.get("/art-profiles/project/{project_id}", response_model=BookArtProfile)
async def get_art_profile_by_project(project_id: str, current_user: UserOut = Depends(get_current_user)):
    """Get the art profile for a project"""
    project = await db.projects.find_one({"id": project_id, "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    profile = await db.book_art_profiles.find_one(
        {"project_id": project_id}, {"_id": 0}
    )
    if not profile:
        raise HTTPException(status_code=404, detail="Art profile not found")
    return profile


@api_router.put("/art-profiles/project/{project_id}", response_model=BookArtProfile)
async def update_art_profile(project_id: str, update: BookArtProfileUpdate, current_user: UserOut = Depends(get_current_user)):
    """Update an existing art profile"""
    project = await db.projects.find_one({"id": project_id, "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    update_data = {k: v for k, v in update.model_dump().items() if v is not None}
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()

    result = await db.book_art_profiles.update_one(
        {"project_id": project_id}, {"$set": update_data}
    )

    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Art profile not found")

    profile = await db.book_art_profiles.find_one(
        {"project_id": project_id}, {"_id": 0}
    )
    return profile


@api_router.post("/ai/art-profile-summary")
async def generate_art_profile_summary(profile: BookArtProfileCreate, current_user: UserOut = Depends(get_current_user)):
    """Generate AI summary for a book art profile with targeted refinement suggestions"""

    context_parts = []
    if profile.genre:
        context_parts.append(f"Genre: {profile.genre}")
    if profile.age_group:
        context_parts.append(f"Age group: {profile.age_group}")
    if profile.mood:
        context_parts.append(f"Mood: {profile.mood}")
    if profile.art_style_preferences:
        context_parts.append(f"Art style preferences: {profile.art_style_preferences}")
    if profile.color_palette:
        context_parts.append(f"Color palette: {profile.color_palette}")
    if profile.reference_notes:
        context_parts.append(f"Reference notes: {profile.reference_notes}")

    user_context = (
        "\n".join(context_parts) if context_parts else "No details provided yet."
    )

    prompt = P.build_art_profile_summary_prompt(user_context=user_context)

    try:
        response = await get_ai_response(BOOK_ART_PROFILE_SYSTEM_PROMPT, prompt)

        # Parse JSON from response
        import json
        import re

        json_match = re.search(r"\{[^{}]*\}", response, re.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group())
            summary = parsed.get("summary", response)
            refinements = parsed.get("refinements", [])
        else:
            summary = response
            refinements = []

        return {
            "summary": summary,
            "refinements": refinements[:3],  # Allow up to 3 refinements
        }

    except Exception as e:
        logger.error(f"Art profile summary generation failed: {e}")
        return {
            "summary": P.FALLBACK_ART_PROFILE_SUMMARY,
            "refinements": P.FALLBACK_ART_PROFILE_REFINEMENTS,
        }


# Scene extraction request model
class SceneExtractionRequest(BaseModel):
    chapter_content: str
    art_profile: Optional[dict] = None


# Scene art prompt request model
class SceneArtPromptRequest(BaseModel):
    project_id: str
    chapter_id: Optional[str] = None
    scene_text: str
    prompt_type: str = "spot_illustration"
    style_preset: str = ""
    art_profile: Optional[dict] = None


SCENE_ART_PROMPT_SYSTEM = P.SCENE_ART_PROMPT_SYSTEM


@api_router.post("/ai/extract-scene")
async def extract_visually_rich_scene(request: SceneExtractionRequest, current_user: UserOut = Depends(get_current_user)):
    """Extract the most visually rich moment from chapter content"""

    # Clean HTML tags from content
    import re

    clean_content = re.sub(r"<[^>]+>", "", request.chapter_content)

    # Truncate for processing
    content_preview = clean_content[:4000]
    if len(clean_content) > 4000:
        content_preview += "... [truncated]"

    prompt = P.build_extract_scene_prompt(content_preview=content_preview)

    try:
        response = await get_ai_response(GLOBAL_SYSTEM_PROMPT, prompt)
        return {"scene": response.strip()}
    except Exception as e:
        logger.error(f"Scene extraction failed: {e}")
        # Return first 500 chars as fallback
        return {"scene": clean_content[:500]}


@api_router.post("/ai/scene-art-prompt")
async def generate_scene_art_prompt(request: SceneArtPromptRequest, current_user: UserOut = Depends(get_current_user)):
    """Generate a structured art prompt from scene text using the Book Art Profile"""

    # Build art profile context
    profile_context = ""
    if request.art_profile:
        profile_parts = []
        if request.art_profile.get("genre"):
            profile_parts.append(f"Genre: {request.art_profile['genre']}")
        if request.art_profile.get("age_group"):
            profile_parts.append(f"Age group: {request.art_profile['age_group']}")
        if request.art_profile.get("mood"):
            profile_parts.append(f"Mood: {request.art_profile['mood']}")
        if request.art_profile.get("art_style_preferences"):
            profile_parts.append(
                f"Art style: {request.art_profile['art_style_preferences']}"
            )
        if request.art_profile.get("color_palette"):
            profile_parts.append(
                f"Color palette: {request.art_profile['color_palette']}"
            )
        if request.art_profile.get("reference_notes"):
            profile_parts.append(
                f"References: {request.art_profile['reference_notes']}"
            )
        if request.art_profile.get("ai_summary"):
            profile_parts.append(
                f"Visual identity: {request.art_profile['ai_summary']}"
            )
        profile_context = "\n".join(profile_parts)

    prompt_type_labels = {
        "cover": "book cover",
        "chapter_header": "chapter header illustration",
        "spot_illustration": "spot illustration",
    }

    prompt_label = prompt_type_labels.get(request.prompt_type, "illustration")

    prompt = P.build_scene_art_prompt(
        profile_context=profile_context,
        style_preset=request.style_preset,
        prompt_label=prompt_label,
        scene_text=request.scene_text,
    )
    try:
        response = await get_ai_response(SCENE_ART_PROMPT_SYSTEM, prompt)

        # Parse JSON from response
        import json
        import re

        # Try to find JSON in response
        json_match = re.search(r"\{[\s\S]*\}", response)
        if json_match:
            try:
                parsed = json.loads(json_match.group())
                return {
                    "main_prompt": parsed.get("main_prompt", response),
                    "refinement_suggestions": parsed.get("refinement_suggestions", [])[
                        :2
                    ],
                    "focus_elements": parsed.get(
                        "focus_elements",
                        {
                            "characters": [],
                            "setting": "Not specified",
                            "action": "Not specified",
                        },
                    ),
                    "response": parsed.get(
                        "main_prompt", response
                    ),  # Legacy compatibility
                }
            except json.JSONDecodeError:
                pass

       # Fallback if JSON parsing fails
        return {
            "main_prompt": response,
            "refinement_suggestions": P.FALLBACK_SCENE_PROMPT_REFINEMENTS,
            "focus_elements": P.FALLBACK_SCENE_PROMPT_FOCUS,
            "response": response,
        }

    except Exception as e:
        logger.error(f"Scene art prompt generation failed: {e}")
        return {
            "main_prompt": P.FALLBACK_SCENE_PROMPT_MAIN,
            "refinement_suggestions": P.FALLBACK_SCENE_PROMPT_REFINEMENTS,
            "focus_elements": P.FALLBACK_SCENE_PROMPT_FOCUS,
            "response": P.FALLBACK_SCENE_PROMPT_MAIN,
        }


# ============== AI IMAGE GENERATION ENDPOINT ==============


class ImageGenerationRequest(BaseModel):
    prompt: str
    size: str = "1024x1024"  # 1024x1024, 1024x1536 (portrait), 1536x1024 (landscape)
    project_id: Optional[str] = None
    chapter_id: Optional[str] = None
    image_type: str = "cover"  # cover, chapter_header, spot_illustration


class ImageGenerationResponse(BaseModel):
    success: bool
    image_base64: Optional[str] = None
    message: str
    asset_id: Optional[str] = None


@api_router.post("/ai/generate-image", response_model=ImageGenerationResponse)
async def generate_image_from_prompt(request: ImageGenerationRequest, current_user: UserOut = Depends(get_current_user)):
    """Generate an image from a text prompt using OpenAI's image generation."""

    if not request.prompt or len(request.prompt.strip()) < 10:
        raise HTTPException(
            status_code=400, detail="Prompt must be at least 10 characters"
        )

    if USE_MOCK_AI:
        try:
            import random

            mock_images_dir = ROOT_DIR / "mock_images"
            mock_images = [path for path in mock_images_dir.iterdir() if path.is_file()]

            if not mock_images:
                return ImageGenerationResponse(
                    success=False,
                    image_base64=None,
                    message="No mock images found in backend/mock_images",
                    asset_id=None,
                )

            selected_image = random.choice(mock_images)
            image_bytes = selected_image.read_bytes()
            image_base64 = base64.b64encode(image_bytes).decode("utf-8")

            return ImageGenerationResponse(
                success=True,
                image_base64=image_base64,
                message=f"Mock image generated from {selected_image.name}",
                asset_id=None,
            )
        except Exception as e:
            logger.error(f"Mock image generation failed: {e}")
            return ImageGenerationResponse(
                success=False,
                image_base64=None,
                message=f"Mock image generation failed: {str(e)}",
                asset_id=None,
            )

    try:
        logger.info(
            f"Generating image with prompt: {request.prompt[:100]}... size={request.size}"
        )

        # Use LiteLLM directly to support size parameter
        from litellm import image_generation
        import requests as http_requests

        params = {
            "model": "gpt-image-1",
            "prompt": request.prompt,
            "n": 1,
            "size": request.size,
            "api_key": OPENAI_API_KEY,
        }

        response = image_generation(**params)

        # Convert response to bytes
        image_bytes = None
        if response.data and len(response.data) > 0:
            img = response.data[0]
            if hasattr(img, "b64_json") and img.b64_json:
                image_bytes = base64.b64decode(img.b64_json)
            elif hasattr(img, "url") and img.url:
                img_response = http_requests.get(img.url)
                image_bytes = img_response.content

        if not image_bytes:
            return ImageGenerationResponse(
                success=False, message="No image was generated", image_base64=None
            )

        # Convert to base64
        image_base64 = base64.b64encode(image_bytes).decode("utf-8")

        # Optionally save as art asset
        asset_id = None
        if request.project_id:
            asset_obj = ArtAsset(
                project_id=request.project_id,
                chapter_id=request.chapter_id,
                type=request.image_type,
                style_preset="ai_generated",
                prompt_used=request.prompt,
                status="generated",
                image_reference=f"data:image/png;base64,{image_base64}",  # Store truncated reference
            )
            doc = asset_obj.model_dump()
            await db.art_assets.insert_one(doc)
            asset_id = asset_obj.id

        logger.info(f"Image generated successfully, asset_id={asset_id}")

        return ImageGenerationResponse(
            success=True,
            image_base64=image_base64,
            message="Painted.",
            asset_id=asset_id,
        )

    except Exception as e:
        logger.error(f"Image generation failed: {e}")
        return ImageGenerationResponse(
            success=False,
            message=f"Nothing came back from the model.",
            image_base64=None,
        )


# ============== TONE PROFILE ENDPOINTS ==============


@api_router.get("/tone-profiles/project/{project_id}", response_model=List[ToneProfile])
async def get_tone_profiles_by_project(project_id: str, current_user: UserOut = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id, "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    profiles = await db.tone_profiles.find(
        {"project_id": project_id}, {"_id": 0}
    ).to_list(1000)
    return profiles


@api_router.get("/tone-profiles/chapter/{chapter_id}", response_model=ToneProfile)
async def get_tone_profile_by_chapter(chapter_id: str, current_user: UserOut = Depends(get_current_user)):
    # Verify chapter belongs to this user via parent project
    chapter = await db.chapters.find_one({"id": chapter_id})
    if not chapter:
        raise HTTPException(status_code=404, detail="Tone profile not found")
    project = await db.projects.find_one({"id": chapter["project_id"], "user_id": current_user.id})
    if not project:
        raise HTTPException(status_code=404, detail="Tone profile not found")
    profile = await db.tone_profiles.find_one({"chapter_id": chapter_id}, {"_id": 0})
    if not profile:
        raise HTTPException(status_code=404, detail="Tone profile not found")
    return profile


# ============== AI ENDPOINTS ==============


@api_router.post("/ai/rewrite", response_model=AIResponse)
@api_router.post("/ai/improve-writing", response_model=AIResponse)
async def rewrite_text(request: RewriteRequest, current_user: UserOut = Depends(get_current_user)):
    if USE_MOCK_AI:
        project = None
        if request.project_id:
            project = await db.projects.find_one({"id": request.project_id}, {"_id": 0})

        response = _mock_rewrite(
            request.content,
            tone_style=(project or {}).get("tone_style", ""),
            pacing_preference=(project or {}).get("pacing_preference", ""),
            target_audience=(project or {}).get("target_audience", ""),
            voice_style=(project or {}).get("voice_style", ""),
            style_notes=(project or {}).get("style_notes", ""),
        )
        return AIResponse(response=response, module="manuscript")

    # Read voice/tone fields from project (if any)
    voice_style = ""
    tone_style = ""
    target_audience = ""
    pacing_preference = ""
    style_notes = ""
    age_group = ""

    project_id = (request.project_id or "").strip()
    if project_id:
        project = await db.projects.find_one({"id": project_id}, {"_id": 0})
        if project:
            voice_style = (project.get("voice_style") or "").strip()
            tone_style = (project.get("tone_style") or "").strip()
            target_audience = (project.get("target_audience") or "").strip()
            pacing_preference = (project.get("pacing_preference") or "").strip()
            style_notes = (project.get("style_notes") or "").strip()
            age_group = (project.get("age_group") or "").strip()

    prompt = P.build_rewrite_prompt(
        content=request.content,
        tone=request.tone,
        voice_style=voice_style,
        tone_style=tone_style,
        target_audience=target_audience,
        pacing_preference=pacing_preference,
        style_notes=style_notes,
        age_group=age_group,
    )

    response = await get_ai_response(MANUSCRIPT_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="manuscript")


@api_router.post("/ai/summarize", response_model=AIResponse)
async def summarize_chapter(request: SummarizeRequest, current_user: UserOut = Depends(get_current_user)):
    prompt = P.build_summarize_prompt(content=request.content)

    response = await get_ai_response(MANUSCRIPT_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="manuscript")


@api_router.post("/ai/analyze-tone-basic", response_model=AIResponse)
async def analyze_tone_basic(request: AnalyzeToneRequest, current_user: UserOut = Depends(get_current_user)):
    if USE_MOCK_AI:
        response = (
            "The voice is conversational and character-driven, with a suspenseful "
            "thread that lets up here and there for a quieter or wryer beat. Reads "
            "middle-grade, somewhere around 4th to 6th grade. Pacing moves a touch "
            "faster in dialogue and slows in the description.\n\n"
            "Two notes: vary sentence length for rhythm, and bring in a little "
            "sensory detail in the quieter stretches — the eye wants something to "
            "rest on."
        )
        return AIResponse(response=response, module="tone")

    prompt = P.build_analyze_tone_basic_prompt(content=request.content)

    response = await get_ai_response(MANUSCRIPT_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="tone")


@api_router.post("/ai/outline", response_model=AIResponse)
async def generate_outline(request: OutlineRequest, current_user: UserOut = Depends(get_current_user)):
    prompt = P.build_outline_prompt(
        project_title=request.project_title or "",
        project_summary=request.project_summary or "",
        chapter_count=request.target_chapter_count,
    )

    response = await get_ai_response(MANUSCRIPT_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="manuscript")


@api_router.post("/ai/workflow-analysis", response_model=AIResponse)
async def analyze_workflow(request: WorkflowAnalysisRequest, current_user: UserOut = Depends(get_current_user)):
    prompt = P.build_workflow_analysis_prompt(status_description=request.status_description)

    response = await get_ai_response(WORKFLOW_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="workflow")


@api_router.post("/ai/workflow-stage", response_model=WorkflowStageAnalysisResponse)
async def analyze_workflow_stage(request: WorkflowStageAnalysisRequest, current_user: UserOut = Depends(get_current_user)):
    """Analyze manuscript and determine workflow stage with next steps"""

    # Build context from request
    context_parts = []

    if request.manuscript:
        # Truncate manuscript for analysis (first 3000 chars for efficiency)
        manuscript_preview = request.manuscript[:3000]
        if len(request.manuscript) > 3000:
            manuscript_preview += "... [truncated for analysis]"
        context_parts.append(f"Current manuscript text:\n{manuscript_preview}")

    if request.section_info:
        context_parts.append(f"Section metadata: {request.section_info}")

    if request.workflow_stage:
        context_parts.append(
            f"Current workflow stage (user-set): {request.workflow_stage}"
        )

    if request.goals:
        context_parts.append(f"Writing goals: {request.goals}")

    if request.time_away:
        context_parts.append(f"Time since last session: {request.time_away}")

    if request.age_group:
        context_parts.append(f"Target age group: {request.age_group}")

    user_context = (
        "\n".join(context_parts)
        if context_parts
        else "No manuscript content provided yet."
    )

    prompt = P.build_workflow_stage_prompt(user_context=user_context)

    try:
        response = await get_ai_response(WORKFLOW_STAGE_SYSTEM_PROMPT, prompt)

        # Parse JSON from response
        import json
        import re

        # Try to extract JSON from response
        json_match = re.search(r"\{[^{}]*\}", response, re.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group())
            stage = parsed.get("stage", "Draft")
            message = parsed.get("message", response)
            next_steps = parsed.get(
                "next_steps", ["Continue writing", "Review your progress"]
            )
        else:
            # Fallback if JSON parsing fails
            stage = "Draft"
            message = response
            next_steps = ["Continue writing", "Review your progress"]

        # Map stage to progress percentage
        stage_progress = {
            "Idea Drop": 10,
            "Outline": 25,
            "Draft": 50,
            "Revise": 70,
            "Polish": 90,
            "Complete": 100,
        }

        progress = stage_progress.get(stage, 50)

        return WorkflowStageAnalysisResponse(
            stage=stage,
            message=message,
            next_steps=next_steps[:2],  # Limit to 2 steps
            progress_percent=progress,
        )

    except Exception as e:
        logger.error(f"Workflow stage analysis failed: {e}")
        # Return sensible defaults
        return WorkflowStageAnalysisResponse(
            stage=P.FALLBACK_WORKFLOW_STAGE,
            message=P.FALLBACK_WORKFLOW_MESSAGE,
            next_steps=P.FALLBACK_WORKFLOW_NEXT_STEPS,
            progress_percent=50,
        )


@api_router.post("/ai/writing-momentum", response_model=WritingMomentumResponse)
async def analyze_writing_momentum(request: WritingMomentumRequest, current_user: UserOut = Depends(get_current_user)):
    """Analyze writing momentum and provide encouragement"""

    # Build context from request
    context_parts = []
    context_parts.append(f"Daily word count: {request.daily_words}")
    context_parts.append(f"Weekly word count: {request.weekly_words}")
    context_parts.append(f"Streak length: {request.streak} days")
    context_parts.append(f"Total manuscript words: {request.total_words}")
    context_parts.append(f"Session duration: {request.session_minutes} minutes")

    if request.time_away:
        context_parts.append(f"Time since last session: {request.time_away}")

    if request.goals:
        context_parts.append(f"Writing goals: {request.goals}")

    if request.age_group:
        context_parts.append(f"Target age group: {request.age_group}")

    user_context = "\n".join(context_parts)

    prompt = P.build_writing_momentum_prompt(user_context=user_context)

    try:
        response = await get_ai_response(WRITING_MOMENTUM_SYSTEM_PROMPT, prompt)

        # Parse JSON from response
        import json
        import re

        # Try to extract JSON from response
        json_match = re.search(r"\{[^{}]*\}", response, re.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group())
            message = parsed.get("message", response)
            suggestions = parsed.get("suggestions", P.FALLBACK_MOMENTUM_SUGGESTIONS)
        else:
            # Fallback if JSON parsing fails
            message = response
            suggestions = P.FALLBACK_MOMENTUM_SUGGESTIONS

        return WritingMomentumResponse(
            message=message, suggestions=suggestions[:2]  # Limit to 2 suggestions
        )

    except Exception as e:
        logger.error(f"Writing momentum analysis failed: {e}")
        return WritingMomentumResponse(
            message=P.momentum_fallback_message(request.streak, request.daily_words),
            suggestions=P.FALLBACK_MOMENTUM_SUGGESTIONS,
        )

        return WritingMomentumResponse(
            message=message,
            suggestions=[
                "Write for just 10 minutes today",
                "Review your last paragraph to get back in the flow",
            ],
        )


@api_router.post("/ai/analyze-tone", response_model=ToneStyleAnalysisResponse)
async def analyze_tone(request: ToneAnalysisRequest, current_user: UserOut = Depends(get_current_user)):
    """Enhanced Tone & Style Analysis with structured output"""

    # Build context from request
    context_parts = []

    if request.content:
        # Truncate content for analysis (first 4000 chars for efficiency)
        content_preview = request.content[:4000]
        if len(request.content) > 4000:
            content_preview += "... [truncated for analysis]"
        context_parts.append(f"Manuscript text:\n{content_preview}")

    if request.section_info:
        context_parts.append(f"Section metadata: {request.section_info}")

    if request.intended_tone:
        context_parts.append(f"Intended tone: {request.intended_tone}")

    if request.goals:
        context_parts.append(f"Writing goals: {request.goals}")

    if request.age_group:
        context_parts.append(f"Target age group: {request.age_group}")

    user_context = "\n".join(context_parts) if context_parts else "No content provided."

    prompt = P.build_analyze_tone_prompt(user_context=user_context)

    try:
        response = await get_ai_response(TONE_STYLE_ANALYSIS_SYSTEM_PROMPT, prompt)

        # Parse JSON from response
        import json
        import re

        # Try to extract JSON from response
        json_match = re.search(r"\{[^{}]*\}", response, re.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group())
            tone_analysis = parsed.get("tone_analysis", "Unable to analyze tone.")
            style_analysis = parsed.get("style_analysis", "Unable to analyze style.")
            suggestions = parsed.get("suggestions", P.FALLBACK_TONE_SUGGESTIONS)
            reading_level = parsed.get("reading_level", P.FALLBACK_READING_LEVEL)
        else:
            # Fallback if JSON parsing fails
            tone_analysis = response[:500] if response else "Unable to analyze tone."
            style_analysis = "See tone analysis for style notes."
            suggestions = P.FALLBACK_TONE_SUGGESTIONS
            reading_level = P.FALLBACK_READING_LEVEL

        # Save tone profile to database
        tone_profile = ToneProfile(
            project_id=request.project_id,
            chapter_id=request.chapter_id,
            detected_tone=tone_analysis,
            reading_level=reading_level,
            pacing_notes=style_analysis,
            voice_notes="",
            suggestions=suggestions,
        )

        # Update or insert tone profile
        if request.chapter_id:
            await db.tone_profiles.update_one(
                {"chapter_id": request.chapter_id},
                {"$set": tone_profile.model_dump()},
                upsert=True,
            )
        else:
            await db.tone_profiles.insert_one(tone_profile.model_dump())

        return ToneStyleAnalysisResponse(
            tone_analysis=tone_analysis,
            style_analysis=style_analysis,
            suggestions=suggestions[:2],  # Limit to 2 suggestions
            reading_level=reading_level,
        )

    except Exception as e:
        logger.error(f"Tone & Style analysis failed: {e}")
        # Return sensible defaults
        return ToneStyleAnalysisResponse(
            tone_analysis=P.FALLBACK_TONE_ANALYSIS,
            style_analysis=P.FALLBACK_STYLE_ANALYSIS,
            suggestions=P.FALLBACK_TONE_SUGGESTIONS,
            reading_level=P.FALLBACK_READING_LEVEL,
        )


@api_router.post("/ai/art-prompts", response_model=AIResponse)
async def generate_art_prompts(request: ArtPromptRequest, current_user: UserOut = Depends(get_current_user)):
    # Get project info
    project = await db.projects.find_one({"id": request.project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if request.prompt_type == "chapter_header":
        prompt = P.build_art_prompt_chapter_header(
            style_preset=request.style_preset,
            context=request.context,
        )
    elif request.prompt_type == "spot_illustration":
        prompt = P.build_art_prompt_spot_illustration(
            style_preset=request.style_preset,
            context=request.context,
        )
    else:  # cover (default)
        prompt = P.build_art_prompt_cover(
            style_preset=request.style_preset,
            project_title=project.get("title", "Untitled"),
            series_name=project.get("series_name") or "",
            context=request.context,
        )
    response = await get_ai_response(ART_STUDIO_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="art")


@api_router.post("/ai/ask-thad", response_model=AIResponse)
async def ask_thad(request: AskThadRequest, current_user: UserOut = Depends(get_current_user)):
    prompt = P.build_ask_thad_prompt(query=request.query, context=request.context or "")

    response = await get_ai_response(GLOBAL_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="general")


@api_router.post("/ai/chat", response_model=AIResponse)
async def chat_with_thad(request: AIChatRequest, current_user: UserOut = Depends(get_current_user)):
    clean_chapter_content = re.sub(
        r"<[^>]+>", "", request.chapter_content or ""
    ).strip()
    selected_text = (request.selected_text or "").strip()

    chapter_excerpt = clean_chapter_content[:6000]
    if len(clean_chapter_content) > 6000:
        chapter_excerpt += "\n\n[Chapter content truncated for chat context]"

    project_context = "No project selected."
    if request.project_id:
        project = await db.projects.find_one({"id": request.project_id}, {"_id": 0})
        if project:
            project_context = (
                f"Project title: {project.get('title', 'Untitled')}\n"
                f"Genre: {project.get('genre') or 'Unknown'}\n"
                f"Age group: {project.get('age_group') or 'Unknown'}\n"
                f"Summary: {project.get('summary') or 'None provided'}"
            )

    if USE_MOCK_AI:
        prompt_lower = request.message.lower()
        focus_text = selected_text or chapter_excerpt
        focus_preview = (
            " ".join(focus_text.split())[:280] or "No chapter content was provided."
        )

        if "summary" in prompt_lower or "summarize" in prompt_lower:
            response = (
                "[mock]: Here is a quick chapter-focused summary.\n\n"
                f"- Main focus: {focus_preview}\n"
                "- Likely next step: tighten the emotional turn or clarify the key beat.\n"
                "- Safe note: this is only a simulated response for development."
            )
        elif "character" in prompt_lower:
            response = (
                "[mock]: Based on the current chapter context, I'd examine how the character's goal, "
                "emotion, and decision are showing up on the page.\n\n"
                f"Relevant excerpt: {focus_preview}\n\n"
                "A good next question would be: What does the character want in this moment, and is that visible in the prose?"
            )
        elif selected_text:
            response = (
                "[mock]: I can see you've highlighted a specific passage.\n\n"
                f"Selected text: {focus_preview}\n\n"
                "If you want, ask about clarity, tone, pacing, or what this passage implies about the scene."
            )
        else:
            response = (
                "[mock]: I reviewed the chapter context you sent.\n\n"
                f"Your question: {request.message}\n"
                f"Chapter preview: {focus_preview}\n\n"
                "I can help with summary, scene understanding, character intent, pacing, tone, or next-step brainstorming."
            )

        return AIResponse(response=response, module="manuscript_chat")

    prompt = P.build_chat_prompt(
        message=request.message,
        project_context=project_context,
        selected_text=selected_text,
        chapter_excerpt=chapter_excerpt,
    )

    response = await get_ai_response(MANUSCRIPT_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="manuscript_chat")


# ============== MARKET INTELLIGENCE ENDPOINTS ==============


@api_router.post("/ai/market/book-ideas", response_model=AIResponse)
async def generate_book_ideas(request: BookIdeasRequest, current_user: UserOut = Depends(get_current_user)):
    prompt = P.build_market_book_ideas_prompt(
        universe=request.universe,
        count=request.count,
    )

    response = await get_ai_response(MARKET_INTELLIGENCE_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="market_intelligence")


@api_router.post("/ai/market/analysis", response_model=AIResponse)
async def analyze_market(request: MarketAnalysisRequest, current_user: UserOut = Depends(get_current_user)):
    age_context = f" for {request.age_group}" if request.age_group else ""
    prompt = P.build_market_analysis_prompt(
        genre=request.genre,
        age_group=request.age_group or "",
    )

    response = await get_ai_response(MARKET_INTELLIGENCE_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="market_intelligence")


@api_router.post("/ai/market/customer-research", response_model=AIResponse)
async def generate_customer_research(request: CustomerResearchRequest, current_user: UserOut = Depends(get_current_user)):
    prompt = P.build_market_customer_research_prompt(book_idea=request.book_idea)

    response = await get_ai_response(MARKET_INTELLIGENCE_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="market_intelligence")


@api_router.post("/ai/market/outline", response_model=AIResponse)
async def generate_market_outline(request: MarketOutlineRequest, current_user: UserOut = Depends(get_current_user)):
    prompt = P.build_market_outline_prompt(
        book_idea=request.book_idea,
        chapter_count=request.chapter_count,
    )

    response = await get_ai_response(MARKET_INTELLIGENCE_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="market_intelligence")


@api_router.post("/ai/market/manuscript-draft", response_model=AIResponse)
async def generate_manuscript_draft(request: ManuscriptDraftRequest, current_user: UserOut = Depends(get_current_user)):
    prompt = P.build_market_manuscript_draft_prompt(
        book_idea=request.book_idea,
        word_count=request.word_count,
    )

    response = await get_ai_response(MARKET_INTELLIGENCE_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="market_intelligence")


@api_router.post("/ai/market/book-description", response_model=AIResponse)
async def generate_book_description(request: BookDescriptionRequest, current_user: UserOut = Depends(get_current_user)):
    prompt = P.build_market_book_description_prompt(
        book_title=request.book_title,
        book_summary=request.book_summary,
    )

    response = await get_ai_response(MARKET_INTELLIGENCE_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="market_intelligence")


@api_router.post("/ai/market/sales-analysis", response_model=AIResponse)
async def analyze_sales_data(request: SalesAnalysisRequest, current_user: UserOut = Depends(get_current_user)):
    prompt = P.build_market_sales_analysis_prompt(sales_data=request.sales_data)

    response = await get_ai_response(MARKET_INTELLIGENCE_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="market_intelligence")


# ============== THAD ONBOARDING ENDPOINT ==============


class ThadWelcomeRequest(BaseModel):
    user_name: str = "Writer"
    book_title: Optional[str] = None
    age_group: Optional[str] = None
    theme: Optional[str] = None
    device_type: str = "desktop"


class ThadWelcomeResponse(BaseModel):
    message: str
    next_steps: List[str]


THAD_SYSTEM_PROMPT = P.THAD_WELCOME_SYSTEM_PROMPT


@api_router.post("/ai/thad/welcome", response_model=ThadWelcomeResponse)
async def generate_thad_welcome(request: ThadWelcomeRequest, current_user: UserOut = Depends(get_current_user)):
    """Generate a personalized welcome message from Thad"""
    import json

    context_parts = []
    if request.user_name:
        context_parts.append(f"User name: {request.user_name}")
    if request.book_title:
        context_parts.append(f"Book title: {request.book_title}")
    if request.age_group:
        context_parts.append(f"Age group: {request.age_group}")
    if request.theme:
        context_parts.append(f"Theme: {request.theme}")
    context_parts.append(f"Device: {request.device_type}")

    user_context = "\n".join(context_parts)

    prompt = P.build_thad_welcome_prompt(user_context=user_context)

    try:
        response = await get_ai_response(THAD_SYSTEM_PROMPT, prompt)

        # Try to parse as JSON
        try:
            # Clean up the response if it has markdown code blocks
            cleaned = response.strip()
            if cleaned.startswith("```"):
                cleaned = cleaned.split("```")[1]
                if cleaned.startswith("json"):
                    cleaned = cleaned[4:]
            cleaned = cleaned.strip()

            data = json.loads(cleaned)
            return ThadWelcomeResponse(
                message=data.get("message", response),
                next_steps=data.get(
                    "next_steps",
                    [
                        P.FALLBACK_WELCOME_STEPS,
                        P.FALLBACK_WELCOME_STEPS,
                        P.FALLBACK_WELCOME_STEPS,
                    ],
                ),
            )
        except json.JSONDecodeError:
            # If not valid JSON, return the raw response with default next steps
            return ThadWelcomeResponse(
                message=response,
                next_steps=[
                    P.FALLBACK_WELCOME_STEPS,
                    P.FALLBACK_WELCOME_STEPS,
                    P.FALLBACK_WELCOME_STEPS,
                ],
            )
    except Exception as e:
        logger.error(f"Thad welcome generation failed: {e}")
        return ThadWelcomeResponse(
            message=P.FALLBACK_WELCOME_MESSAGE,
            next_steps=P.FALLBACK_WELCOME_STEPS,
        )


# ============== THAD GUIDED TOUR ENDPOINT ==============

TOUR_STEPS = P.TOUR_STEPS


class ThadTourRequest(BaseModel):
    user_name: str = "Writer"
    book_title: Optional[str] = None
    age_group: Optional[str] = None
    theme: Optional[str] = None
    device_type: str = "desktop"
    current_step: int = 0


class ThadTourResponse(BaseModel):
    step_number: int
    total_steps: int
    area: str
    message: str
    is_final: bool
    final_actions: Optional[List[str]] = None


THAD_TOUR_SYSTEM_PROMPT = P.THAD_TOUR_SYSTEM_PROMPT


@api_router.post("/ai/thad/tour", response_model=ThadTourResponse)
async def generate_thad_tour_step(request: ThadTourRequest, current_user: UserOut = Depends(get_current_user)):
    """Generate a guided tour step from Thad"""
    import json

    current_step = min(request.current_step, len(TOUR_STEPS) - 1)
    is_final = current_step >= len(TOUR_STEPS) - 1
    step_info = TOUR_STEPS[current_step]

    # Build context
    context_parts = [f"User name: {request.user_name}"]
    if request.book_title:
        context_parts.append(f"Book title: {request.book_title}")
    if request.age_group:
        context_parts.append(f"Age group: {request.age_group}")
    if request.theme:
        context_parts.append(f"Theme: {request.theme}")
    context_parts.append(f"Device: {request.device_type}")

    user_context = "\n".join(context_parts)

    prompt = P.build_thad_tour_step_prompt(
        area=step_info["area"],
        description=step_info["description"],
        user_context=user_context,
        is_final=is_final,
    )

    try:
        response = await get_ai_response(THAD_TOUR_SYSTEM_PROMPT, prompt)

        # Try to parse JSON
        try:
            cleaned = response.strip()
            if cleaned.startswith("```"):
                cleaned = cleaned.split("```")[1]
                if cleaned.startswith("json"):
                    cleaned = cleaned[4:]
            cleaned = cleaned.strip()

            data = json.loads(cleaned)
            message = data.get("message", response)
        except json.JSONDecodeError:
            message = response

    except Exception as e:
        logger.error(f"Thad tour generation failed: {e}")
        message = step_info["description"]

    final_actions = None
    if is_final:
        final_actions = P.FINAL_TOUR_ACTIONS

    return ThadTourResponse(
        step_number=current_step + 1,
        total_steps=len(TOUR_STEPS),
        area=step_info["area"],
        message=message,
        is_final=is_final,
        final_actions=final_actions,
    )


# ============== IMPORT ANALYSIS ENDPOINTS ==============


@api_router.post("/ai/import/analyze")
async def analyze_imported_manuscript(request: ImportAnalysisRequest, current_user: UserOut = Depends(get_current_user)):
    """Analyze an imported manuscript and provide comprehensive insights"""

    word_count = len(request.content.split())
    project_context = ""
    if request.project_id:
        project = await db.projects.find_one({"id": request.project_id}, {"_id": 0})
        if project:
            if project.get("title"):
                project_context += f"Project Title:\n{project['title']}\n\n"
            if project.get("summary"):
                project_context += f"Project Summary:\n{project['summary']}\n\n"

    content_truncated = request.content[:15000]
    if len(request.content) > 15000:
        content_truncated += "..."

    prompt = P.build_import_analysis_prompt(
        content_truncated=content_truncated,
        word_count=word_count,
        filename=request.filename or "",
        project_context=project_context.strip(),
    )

    analysis_response = await get_ai_response(IMPORT_ANALYSIS_SYSTEM_PROMPT, prompt)

    # Basic detection for structured response
    structure_issues = []
    notes_detected = []
    style_issues = []
    formatting_issues = []
    lore_issues = []

    # Simple detection of common patterns
    import re

    # Detect notes/comments
    note_patterns = [
        r"\[TODO[^\]]*\]",
        r"\[NOTE[^\]]*\]",
        r"\[FIXME[^\]]*\]",
        r"\(Author note:[^)]*\)",
        r"\{\{[^}]*\}\}",
        r"<!--[^>]*-->",
    ]
    for pattern in note_patterns:
        matches = re.findall(pattern, request.content, re.IGNORECASE)
        notes_detected.extend(matches)

    # Estimate reading level based on average sentence and word length
    sentences = re.split(r"[.!?]+", request.content)
    avg_words_per_sentence = word_count / max(len(sentences), 1)

    if avg_words_per_sentence < 10:
        reading_level = "Early Reader (K-2nd grade)"
    elif avg_words_per_sentence < 15:
        reading_level = "Elementary (3rd-5th grade)"
    elif avg_words_per_sentence < 20:
        reading_level = "Middle Grade (6th-8th grade)"
    else:
        reading_level = "Young Adult/Adult"

    # Recommended actions based on analysis
    recommended_actions = ["full_qa"]
    if len(notes_detected) > 0:
        recommended_actions.extend(["remove_notes", "store_notes"])

    # Detect chapter markers to help user understand what will be split
    # Use more flexible patterns that work with various line ending styles
    chapter_patterns = [
        r"(?:^|\n+)\s*(CHAPTER\s+\d+)",
        r"(?:^|\n+)\s*(Chapter\s+\d+)",
        r"(?:^|\n+)\s*(CHAPTER\s+[IVXLCDM]+)",
        r"(?:^|\n+)\s*(Chapter\s+(?:One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|Eleven|Twelve))",
        r"(?:^|\n+)\s*(Prologue)",
        r"(?:^|\n+)\s*(Epilogue)",
        r"(?:^|\n+)\s*(Part\s+\d+)",
    ]
    detected_chapters = []
    for pattern in chapter_patterns:
        matches = re.findall(pattern, request.content, re.IGNORECASE | re.MULTILINE)
        detected_chapters.extend([m.strip() for m in matches])

    # Remove duplicates and sort
    detected_chapters = list(set(detected_chapters))
    chapters_count = len(detected_chapters)

    logger.info(
        f"Import analysis: Detected {chapters_count} chapter markers: {detected_chapters[:10]}"
    )

    if chapters_count > 1:
        recommended_actions.append("split_chapters")
    elif "chapter" in request.content.lower() or "Chapter" in request.content:
        recommended_actions.append("split_chapters")
    recommended_actions.extend(["autoformat", "extract_summaries"])

    return {
        "analysis": analysis_response,
        "structure_issues": structure_issues,
        "notes_detected": notes_detected[:20],  # Limit to first 20
        "style_issues": style_issues,
        "formatting_issues": formatting_issues,
        "lore_issues": lore_issues,
        "word_count": word_count,
        "estimated_reading_level": reading_level,
        "recommended_actions": list(set(recommended_actions)),
        "detected_chapters_count": chapters_count,
        "detected_chapters_preview": detected_chapters[
            :15
        ],  # Show first 15 chapter markers
    }


@api_router.post("/ai/import/action", response_model=AIResponse)
async def execute_import_action(request: ImportActionRequest, current_user: UserOut = Depends(get_current_user)):
    """Execute a specific action on imported manuscript content"""

    action_prompts = P.IMPORT_ACTION_PROMPTS

    if request.action not in action_prompts:
        raise HTTPException(status_code=400, detail=f"Unknown action: {request.action}")

    prompt = action_prompts[request.action].format(content=request.content[:20000])

    response = await get_ai_response(IMPORT_ANALYSIS_SYSTEM_PROMPT, prompt)
    return AIResponse(response=response, module="import_analysis")


# ============== IMPLEMENT IMPORT ACTION ENDPOINT ==============


class ImplementActionRequest(BaseModel):
    action: str  # autoformat, remove_notes, etc.
    original_content: str  # The original chapter content
    chapter_id: Optional[str] = None
    project_id: Optional[str] = None
    response_text: Optional[str] = None
    # For actions that produce modified content
    apply_content: bool = True  # Whether to apply the AI-generated content
    # For notes-related actions
    extracted_notes: Optional[List[str]] = None


class ImplementActionResponse(BaseModel):
    success: bool
    message: str
    action: str
    chapter_updated: bool = False
    notes_created: int = 0
    new_content_preview: Optional[str] = None
    implemented: bool = False
    applied_improvement_id: Optional[str] = None


@api_router.post("/ai/import/implement", response_model=ImplementActionResponse)
async def implement_import_action(request: ImplementActionRequest, current_user: UserOut = Depends(get_current_user)):
    """
    Safely implement an import action by saving the AI result as a structured
    applied improvement record without directly rewriting chapter content.
    """

    try:
        response_text = (request.response_text or "").strip()
        if not response_text and request.extracted_notes:
            response_text = "\n".join(request.extracted_notes[:20])
        if not response_text:
            response_text = "Recorded - nothing to apply yet."

        improvement = AppliedImprovement(
            action=request.action,
            response_text=response_text,
            project_id=request.project_id,
            chapter_id=request.chapter_id,
            original_content_preview=(request.original_content or "")[:500],
            implemented=True,
        )
        await db.applied_improvements.insert_one(improvement.model_dump())

        return ImplementActionResponse(
            success=True,
            message="Saved.",
            action=request.action,
            chapter_updated=False,
            notes_created=0,
            new_content_preview=response_text[:500],
            implemented=True,
            applied_improvement_id=improvement.id,
        )

    except Exception as e:
        logger.error(f"Implement action failed: {e}")
        raise HTTPException(
            status_code=500, detail=f"Failed to implement action: {str(e)}"
        )


# ============== SPLIT CHAPTERS ENDPOINT ==============


class SplitChaptersRequest(BaseModel):
    content: str
    project_id: Optional[str] = None
    manuscript_id: Optional[str] = None


class ChapterInfo(BaseModel):
    chapter_number: int
    title: str
    content: str
    start_position: int = 0


class SplitChaptersResponse(BaseModel):
    success: bool
    chapters_created: int
    chapters: List[dict]
    message: str


def detect_chapters_regex(content: str) -> List[dict]:
    """Regex-based chapter detection - more reliable than AI for structure"""
    import re

    chapters = []
    chapter_positions = []

    # Split content into lines for line-by-line processing
    lines = content.split("\n")
    current_pos = 0

    # Word-to-number mapping for written numbers
    word_numbers = r"One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|Eleven|Twelve|Thirteen|Fourteen|Fifteen|Sixteen|Seventeen|Eighteen|Nineteen|Twenty|Twenty[\s-]?One|Twenty[\s-]?Two|Twenty[\s-]?Three|Twenty[\s-]?Four|Twenty[\s-]?Five|Thirty|Forty|Fifty"

    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            current_pos += len(line) + 1
            continue

        title = None
        matched = False
        chapter_num_str = None

        # Pattern 1: "Chapter X" or "CHAPTER X" with optional title (supports numbers 1-999, Roman numerals, and written numbers)
        match = re.match(
            rf"^(Chapter|CHAPTER)\s+(\d{{1,3}}|[IVXLCDM]+|{word_numbers})[\s:\-–—\.]*(.*)$",
            line_stripped,
            re.IGNORECASE,
        )
        if match:
            matched = True
            chapter_num_str = match.group(2)
            title = match.group(3).strip() if match.group(3) else None

        # Pattern 2: "Part X" with optional title
        if not matched:
            match = re.match(
                rf"^(Part)\s+(\d{{1,3}}|[IVXLCDM]+|{word_numbers})[\s:\-–—\.]*(.*)$",
                line_stripped,
                re.IGNORECASE,
            )
            if match:
                matched = True
                chapter_num_str = match.group(2)
                title = match.group(3).strip() if match.group(3) else None

        # Pattern 3: "Prologue", "Epilogue", "Introduction", "Preface" as chapter markers
        if not matched:
            match = re.match(
                r"^(Prologue|Epilogue|Introduction|Preface|Foreword|Afterword)[\s:\-–—\.]*(.*)$",
                line_stripped,
                re.IGNORECASE,
            )
            if match:
                matched = True
                chapter_num_str = match.group(1)
                title = match.group(2).strip() if match.group(2) else match.group(1)

        # Pattern 4: Standalone number followed by period and title "1. Title"
        if not matched:
            match = re.match(r"^(\d{1,3})\.\s+([A-Z].{2,80})$", line_stripped)
            if match:
                matched = True
                chapter_num_str = match.group(1)
                title = match.group(2).strip()

        if matched:
            # Avoid duplicates (same position)
            if not any(abs(pos - current_pos) < 10 for pos in chapter_positions):
                chapter_positions.append(current_pos)

                # Generate title if not found
                if not title or len(title) < 2:
                    title = f"Chapter {chapter_num_str}"

                # Truncate overly long titles
                if len(title) > 80:
                    title = title[:80].rsplit(" ", 1)[0]

                chapters.append(
                    {"position": current_pos, "marker": line_stripped, "title": title}
                )

        current_pos += len(line) + 1

    # Sort by position (should already be sorted, but just in case)
    chapters.sort(key=lambda x: x["position"])

    # Assign chapter numbers and build final structure
    result = []
    for i, ch in enumerate(chapters):
        result.append(
            {
                "chapter_number": i + 1,
                "title": ch["title"],
                "start_position": ch["position"],
                "marker": ch["marker"],
            }
        )

    return result


@api_router.post("/ai/import/split-chapters", response_model=SplitChaptersResponse)
async def split_and_create_chapters(request: SplitChaptersRequest, current_user: UserOut = Depends(get_current_user)):
    """
    Detect chapter breaks in content and create separate chapter records.
    Uses regex-based detection for reliability, with AI enhancement for titles.
    """
    import re

    if not request.content or len(request.content.strip()) < 100:
        raise HTTPException(
            status_code=400, detail="Content too short to split into chapters"
        )

    if not request.project_id and not request.manuscript_id:
        raise HTTPException(
            status_code=400, detail="Either project_id or manuscript_id is required"
        )

    content = request.content
    logger.info(f"Split chapters: Received {len(content)} characters")
    logger.info(f"Split chapters: First 200 chars: {content[:200]}")
    logger.info(f"Split chapters: Last 200 chars: {content[-200:]}")

    # Use regex-based detection (more reliable for structure)
    chapters_data = detect_chapters_regex(content)
    logger.info(f"Split chapters: Detected {len(chapters_data)} chapter markers")
    for ch in chapters_data:
        logger.info(
            f"  Chapter {ch['chapter_number']}: '{ch['title']}' at position {ch['start_position']}"
        )

    if not chapters_data or len(chapters_data) == 0:
        # No chapters detected - try to use AI as fallback for unstructured content
        logger.info("No chapter markers found, treating as single chapter")
        chapters_data = [
            {
                "chapter_number": 1,
                "title": "Chapter 1",
                "start_position": 0,
                "marker": "",
            }
        ]

    # Now split the content and create chapter records
    created_chapters = []

    for i, chapter_info in enumerate(chapters_data):
        chapter_number = chapter_info.get("chapter_number", i + 1)
        chapter_title = chapter_info.get("title", f"Chapter {chapter_number}")
        start_pos = chapter_info.get("start_position", 0)

        # Find the end position (start of next chapter or end of content)
        if i + 1 < len(chapters_data):
            end_pos = chapters_data[i + 1].get("start_position", len(content))
        else:
            end_pos = len(content)

        # Extract chapter content
        chapter_content = content[start_pos:end_pos].strip()

        # Skip if chapter content is too short (likely a detection error)
        if len(chapter_content) < 20:
            logger.warning(
                f"Skipping chapter {chapter_number}: content too short ({len(chapter_content)} chars)"
            )
            continue

        # Clean up title if it's too long or contains the chapter marker
        if len(chapter_title) > 100:
            chapter_title = chapter_title[:100].rsplit(" ", 1)[0] + "..."

        # Remove leading "Chapter X:" from title if present
        title_clean = re.sub(
            r"^(Chapter\s+\d+|Chapter\s+[IVXLCDM]+|Part\s+\d+)[\s:\-–—]*",
            "",
            chapter_title,
            flags=re.IGNORECASE,
        ).strip()
        if title_clean and len(title_clean) > 2:
            chapter_title = title_clean

        logger.info(
            f"Creating chapter {chapter_number}: '{chapter_title}' ({len(chapter_content)} chars)"
        )

        # Create chapter record
        chapter_obj = Chapter(
            project_id=request.project_id or "",
            manuscript_id=request.manuscript_id,
            chapter_number=chapter_number,
            title=chapter_title,
            content=f"<p>{chapter_content.replace(chr(10), '</p><p>')}</p>",
            status="draft",
        )

        doc = chapter_obj.model_dump()
        await db.chapters.insert_one(doc)

        created_chapters.append(
            {
                "id": chapter_obj.id,
                "chapter_number": chapter_number,
                "title": chapter_title,
                "word_count": len(chapter_content.split()),
            }
        )

    # Update project/manuscript word count
    total_words = sum(c.get("word_count", 0) for c in created_chapters)
    if request.project_id:
        await db.projects.update_one(
            {"id": request.project_id},
            {
                "$inc": {"word_count": total_words},
                "$set": {"updated_at": datetime.now(timezone.utc).isoformat()},
            },
        )

    logger.info(
        f"Split chapters complete: Created {len(created_chapters)} chapters, {total_words} total words"
    )

    return SplitChaptersResponse(
        success=True,
        chapters_created=len(created_chapters),
        chapters=created_chapters,
        message=(
            f"Split into {len(created_chapters)} chapter"
            f"{'s' if len(created_chapters) != 1 else ''}."
        ),
    )


# ============== EXPORT ENDPOINTS ==============


class ExportRequest(BaseModel):
    project_id: str
    include_title_page: bool = True
    include_chapter_numbers: bool = True
    author_override: Optional[str] = None  # if blank, falls back to user.display_name


def strip_html_tags(html_content: str) -> str:
    """Remove HTML tags and decode entities"""
    if not html_content:
        return ""
    # Remove HTML tags
    text = re.sub(r"<br\s*/?>", "\n", html_content)
    text = re.sub(r"</p>", "\n\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    # Decode common entities
    text = text.replace("&nbsp;", " ")
    text = text.replace("&amp;", "&")
    text = text.replace("&lt;", "<")
    text = text.replace("&gt;", ">")
    text = text.replace("&quot;", '"')
    text = text.replace("&#39;", "'")
    # Clean up extra whitespace
    text = re.sub(r"\n\s*\n\s*\n", "\n\n", text)
    return text.strip()


def sanitize_for_pdf(text: str) -> str:
    """Sanitize text for PDF output with standard fonts"""
    if not text:
        return ""
    # Replace common Unicode characters with ASCII equivalents
    replacements = {
        "\u2018": "'",  # Left single quotation mark
        "\u2019": "'",  # Right single quotation mark
        "\u201c": '"',  # Left double quotation mark
        "\u201d": '"',  # Right double quotation mark
        "\u2013": "-",  # En dash
        "\u2014": "--",  # Em dash
        "\u2026": "...",  # Ellipsis
        "\u00a0": " ",  # Non-breaking space
        "\u2022": "*",  # Bullet
        "\u00b7": "*",  # Middle dot
    }
    for unicode_char, ascii_char in replacements.items():
        text = text.replace(unicode_char, ascii_char)

    # Encode to latin-1, replacing unsupported characters
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text

async def _resolve_export_author(
    request: ExportRequest, current_user: UserOut
) -> str:
    """Pick the author name for an export: override → display_name → email → ''."""
    candidate = (request.author_override or "").strip()
    if candidate:
        return candidate
    if current_user.display_name:
        return current_user.display_name
    if current_user.email:
        # Just the local part, capitalised — better than blank.
        return current_user.email.split("@")[0].replace(".", " ").title()
    return ""

async def _fetch_cover_for_project(project_id: str) -> tuple[Optional[bytes], str]:
    """
    Find the most recent cover ArtAsset for a project and return its bytes.

    Returns (bytes, media_type). If no cover exists or the bytes can't be
    decoded, returns (None, "").

    Art assets store images as data URLs in `image_reference`. We decode the
    base64 payload back into raw bytes here so the EPUB can embed it.
    """
    cursor = (
        db.art_assets.find(
            {"project_id": project_id, "type": "cover"}, {"_id": 0}
        )
        .sort("created_at", -1)
    )
    async for asset in cursor:
        ref = asset.get("image_reference") or ""
        # The ManuscriptWorkspace stores covers as `data:image/png;base64,XXXX...`
        # but the comment in server.py says only the first 100 chars are stored
        # as a reference. We try to decode whatever's there; if it's not a full
        # data URL, skip and try the next asset.
        if not ref.startswith("data:"):
            continue
        try:
            header, payload = ref.split(",", 1)
            # header looks like "data:image/png;base64"
            media_type = header.split(";")[0].removeprefix("data:")
            image_bytes = base64.b64decode(payload)
            if len(image_bytes) > 64:  # crude sanity: ignore truncated stubs
                return image_bytes, media_type or "image/png"
        except Exception:
            logger.warning("Cover decode failed for asset %s", asset.get("id"))
            continue
    return None, ""

@api_router.post("/export/docx")
async def export_to_docx(request: ExportRequest, current_user: UserOut = Depends(get_current_user)):
    """Export a project (all chapters) to DOCX format"""

    # Get project — scoped to current user
    project = await db.projects.find_one({"id": request.project_id, "user_id": current_user.id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Get all chapters sorted by chapter number
    chapters = (
        await db.chapters.find({"project_id": request.project_id}, {"_id": 0})
        .sort("chapter_number", 1)
        .to_list(1000)
    )

    # Create document
    doc = DocxDocument()

    # Add title page
    if request.include_title_page:
        title = doc.add_heading(project.get("title", "Untitled"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if project.get("series_name"):
            series = doc.add_paragraph(f"Part of: {project['series_name']}")
            series.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if project.get("summary"):
            doc.add_paragraph()
            summary = doc.add_paragraph(project["summary"])
            summary.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    # Add chapters
    for chapter in chapters:
        # Chapter heading
        chapter_title = chapter.get(
            "title", f"Chapter {chapter.get('chapter_number', '?')}"
        )
        if request.include_chapter_numbers:
            heading_text = (
                f"Chapter {chapter.get('chapter_number', '?')}: {chapter_title}"
            )
        else:
            heading_text = chapter_title

        doc.add_heading(heading_text, 1)

        # Chapter content
        content = strip_html_tags(chapter.get("content", ""))

        # Split into paragraphs and add
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

        doc.add_page_break()

    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Generate filename
    safe_title = re.sub(r"[^\w\s-]", "", project.get("title", "export")).strip()[:50]
    filename = f"{safe_title}.docx"

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


class CustomPDF(FPDF):
    """Custom PDF class with header/footer support"""

    def __init__(self, title=""):
        super().__init__()
        self.document_title = title

    def header(self):
        if self.page_no() > 1:  # Skip header on title page
            self.set_font("Helvetica", "I", 9)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, self.document_title, 0, 0, "C")
            self.ln(15)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 9)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Page {self.page_no()}", 0, 0, "C")


@api_router.post("/export/pdf")
async def export_to_pdf(request: ExportRequest, current_user: UserOut = Depends(get_current_user)):
    """Export a project (all chapters) to PDF format"""

    # Get project — scoped to current user
    project = await db.projects.find_one({"id": request.project_id, "user_id": current_user.id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Get all chapters sorted by chapter number
    chapters = (
        await db.chapters.find({"project_id": request.project_id}, {"_id": 0})
        .sort("chapter_number", 1)
        .to_list(1000)
    )

    # Create PDF
    title = sanitize_for_pdf(project.get("title", "Untitled"))
    pdf = CustomPDF(title)
    pdf.set_auto_page_break(auto=True, margin=25)

    # Add title page
    if request.include_title_page:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 28)
        pdf.ln(60)
        pdf.multi_cell(0, 12, title, align="C")

        if project.get("series_name"):
            pdf.set_font("Helvetica", "I", 14)
            pdf.ln(10)
            series_text = sanitize_for_pdf(f"Part of: {project['series_name']}")
            pdf.multi_cell(0, 8, series_text, align="C")

        if project.get("summary"):
            pdf.set_font("Helvetica", "", 11)
            pdf.ln(20)
            summary_text = sanitize_for_pdf(project["summary"])
            pdf.multi_cell(0, 6, summary_text, align="C")

    # Add chapters
    for chapter in chapters:
        pdf.add_page()

        # Chapter heading
        chapter_title = chapter.get(
            "title", f"Chapter {chapter.get('chapter_number', '?')}"
        )
        if request.include_chapter_numbers:
            heading_text = (
                f"Chapter {chapter.get('chapter_number', '?')}: {chapter_title}"
            )
        else:
            heading_text = chapter_title

        heading_text = sanitize_for_pdf(heading_text)
        pdf.set_font("Helvetica", "B", 18)
        pdf.multi_cell(0, 10, heading_text)
        pdf.ln(5)

        # Chapter content
        content = strip_html_tags(chapter.get("content", ""))
        content = sanitize_for_pdf(content)

        pdf.set_font("Helvetica", "", 11)

        # Split into paragraphs
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            if para_text.strip():
                pdf.multi_cell(0, 6, para_text.strip())
                pdf.ln(3)

    # Output to bytes
    pdf_bytes = pdf.output()
    file_stream = io.BytesIO(pdf_bytes)
    file_stream.seek(0)

    # Generate filename
    safe_title = re.sub(r"[^\w\s-]", "", project.get("title", "export")).strip()[:50]
    filename = f"{safe_title}.pdf"

    return StreamingResponse(
        file_stream,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

@api_router.post("/export/markdown")
async def export_to_markdown(
    request: ExportRequest, current_user: UserOut = Depends(get_current_user)
):
    """Export a project to a Markdown file."""

    project = await db.projects.find_one(
        {"id": request.project_id, "user_id": current_user.id}, {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    chapters = (
        await db.chapters.find({"project_id": request.project_id}, {"_id": 0})
        .sort("chapter_number", 1)
        .to_list(1000)
    )

    author = await _resolve_export_author(request, current_user)

    md_bytes = E.build_markdown(
        project=project,
        chapters=chapters,
        author=author,
        include_title_page=request.include_title_page,
        include_chapter_numbers=request.include_chapter_numbers,
    )

    filename = E.safe_filename(project.get("title", "export"), "md")
    return StreamingResponse(
        io.BytesIO(md_bytes),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@api_router.post("/export/epub")
async def export_to_epub(
    request: ExportRequest, current_user: UserOut = Depends(get_current_user)
):
    """Export a project to an EPUB 3 file."""

    project = await db.projects.find_one(
        {"id": request.project_id, "user_id": current_user.id}, {"_id": 0}
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    chapters = (
        await db.chapters.find({"project_id": request.project_id}, {"_id": 0})
        .sort("chapter_number", 1)
        .to_list(1000)
    )

    author = await _resolve_export_author(request, current_user)
    cover_bytes, cover_media_type = await _fetch_cover_for_project(
        request.project_id
    )

    epub_bytes = E.build_epub(
        project=project,
        chapters=chapters,
        author=author,
        include_title_page=request.include_title_page,
        include_chapter_numbers=request.include_chapter_numbers,
        cover_bytes=cover_bytes,
        cover_media_type=cover_media_type or "image/jpeg",
    )

    filename = E.safe_filename(project.get("title", "export"), "epub")
    return StreamingResponse(
        io.BytesIO(epub_bytes),
        media_type="application/epub+zip",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

# ── Status endpoints ──────────────────────────────────────────────────────────
@api_router.get("/")
async def root():
    return {"message": "Publish Itt API is running"}


@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "service": "Publish Itt"}


# ── Wire everything into the app ──────────────────────────────────────────────

# Give auth router access to the database
auth_set_db(db)

# Include routers
app.include_router(api_router)
app.include_router(auth_router)

# CORS — reads from CORS_ORIGINS env var; falls back to localhost:3000 for safety
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "http://localhost:3000").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def create_indexes():
    """Create indexes for user_id scoping — safe to run multiple times."""
    await db.projects.create_index([("user_id", 1)])
    await db.projects.create_index([("user_id", 1), ("id", 1)])
    await db.writing_sessions.create_index([("user_id", 1)])
    await db.writing_sessions.create_index([("user_id", 1), ("date", 1)])
    await db.users.create_index([("email", 1)], unique=True)
    await db.notes.create_index([("user_id", 1)])
    await db.notes.create_index([("user_id", 1), ("parent_id", 1)])
    await db.versions.create_index([("user_id", 1)])
    await db.versions.create_index([("user_id", 1), ("parent_id", 1)])
    await db.style_presets.create_index([("user_id", 1)])
    logger.info("MongoDB indexes created/verified.")


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
